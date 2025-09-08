#!/usr/bin/env python
"""
Test the integrated templated document system
"""

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

from api.models import Trip
from utils.docgen.trip_document_generator import TripDocumentGenerator

try:
    from docx import Document as DocxDocument
    print("python-docx is available")
except ImportError:
    print("python-docx not available")
    exit(1)

def test_integrated_system():
    """Test the integrated templated document system"""
    
    # Get first available trip
    trip = Trip.objects.first()
    if not trip:
        print("No trips available for testing")
        return
    
    print(f"Testing integrated templated system for Trip {trip.trip_number}")
    print(f"Trip type: {trip.type}, Has patient: {trip.patient is not None}")
    print()
    
    # Create document generator
    generator = TripDocumentGenerator(str(trip.id))
    
    # Test 1: Check template path resolution
    print("=== Testing Template Path Resolution ===")
    test_templates = ["CustomerItin.docx", "HandlingRequest.docx", "PaymentAgreement.docx"]
    
    for template_name in test_templates:
        template_path = generator._get_template_path(template_name)
        is_templated = "templated" in str(template_path)
        exists = template_path.exists()
        print(f"  {template_name}: {'✓ Templated' if is_templated else '✗ Original'} | {'Exists' if exists else 'Missing'}")
        print(f"    Path: {template_path}")
    print()
    
    # Test 2: Check placeholder mappings include lowercase variants
    print("=== Testing Enhanced Placeholder Mappings ===")
    try:
        trip_data = generator._prepare_trip_data()
        patient_data = generator._prepare_patient_data() if trip.patient else None
        placeholders = generator._create_placeholder_mappings(trip_data, None, patient_data)
        
        # Check for both uppercase and lowercase variants
        test_pairs = [
            ('{{TRIP_NUMBER}}', '{{trip_number}}'),
            ('{{PATIENT_NAME}}', '{{patient_name}}'),
            ('{{AIRCRAFT_TAIL}}', '{{aircraft_tail}}'),
            ('{{PRIMARY_PILOT}}', '{{primary_pilot}}'),
        ]
        
        for uppercase, lowercase in test_pairs:
            upper_val = placeholders.get(uppercase, 'MISSING')
            lower_val = placeholders.get(lowercase, 'MISSING')
            match = upper_val == lower_val if upper_val != 'MISSING' and lower_val != 'MISSING' else False
            print(f"  {uppercase} / {lowercase}: {'✓ Match' if match else '✗ Mismatch'}")
            if not match:
                print(f"    Upper: {upper_val}")
                print(f"    Lower: {lower_val}")
        
        special_lowercase = placeholders.get('{{trip_itinerary}}', 'MISSING')
        print(f"  {{trip_itinerary}}: {special_lowercase}")
        print()
        
    except Exception as e:
        print(f"Error testing placeholders: {e}")
        return
    
    # Test 3: Generate customer itinerary (should use templated version)
    print("=== Testing Customer Itinerary Generation ===")
    try:
        customer_doc = generator._generate_customer_itinerary()
        if customer_doc:
            print(f"✓ Generated: {customer_doc.filename}")
            
            # Verify it contains trip data
            doc_path = Path(customer_doc.file_path)
            if doc_path.exists():
                doc = DocxDocument(doc_path)
                all_text = ""
                
                for paragraph in doc.paragraphs:
                    all_text += paragraph.text + "\n"
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            all_text += cell.text + " "
                
                # Check for data population
                test_data = [
                    trip.trip_number,
                    placeholders.get('{{PATIENT_NAME}}', ''),
                    placeholders.get('{{DEPARTURE_DATE}}', '')
                ]
                
                found_count = 0
                for data in test_data:
                    if data and data != 'N/A' and data in all_text:
                        found_count += 1
                
                print(f"  Data population: {found_count}/{len([d for d in test_data if d])} fields found")
                
                # Check for unreplaced placeholders
                unreplaced_count = len([p for p in ['{{patient_name}}', '{{trip_number}}', '{{trip_date}}'] if p in all_text])
                print(f"  Unreplaced placeholders: {unreplaced_count}")
                
                if found_count > 0 and unreplaced_count == 0:
                    print(f"  🎉 SUCCESS: Document properly populated with trip data!")
                else:
                    print(f"  ⚠️  Issues found with data population")
            else:
                print(f"  ✗ Generated file not found: {doc_path}")
        else:
            print(f"  ✗ Customer itinerary generation failed")
    except Exception as e:
        print(f"  Error generating customer itinerary: {e}")
    print()
    
    # Test 4: Test other document types
    print("=== Testing Other Document Types ===")
    other_tests = [
        ('handling_request', generator._generate_handling_request),
        ('payment_agreement', generator._generate_payment_agreement),
    ]
    
    for doc_type, method in other_tests:
        try:
            doc = method()
            if doc:
                print(f"  ✓ {doc_type}: {doc.filename}")
            else:
                print(f"  ✗ {doc_type}: Generation failed")
        except Exception as e:
            print(f"  ✗ {doc_type}: Error - {e}")
    
    print("\n=== Integration Test Complete ===")

if __name__ == "__main__":
    test_integrated_system()