#!/usr/bin/env python
"""
Script to check the new templated CustomerItin document for placeholders
"""

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

try:
    from docx import Document as DocxDocument
    print("python-docx is available")
except ImportError:
    print("python-docx not available")
    exit(1)

def check_templated_document():
    """Check the templated CustomerItin document for placeholders"""
    template_path = Path("utils/docgen/documents/templated/CustomerItin.docx")
    
    if not template_path.exists():
        print(f"Templated document not found: {template_path}")
        return
    
    print(f"Checking templated document: {template_path}")
    
    try:
        doc = DocxDocument(template_path)
        all_text = ""
        
        # Extract all text from document
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only count non-empty paragraphs
                all_text += paragraph.text + "\n"
                paragraph_count += 1
        
        # Check for tables too
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        print(f"Document stats:")
        print(f"  - Total characters: {len(all_text)}")
        print(f"  - Paragraphs with content: {paragraph_count}")
        print(f"  - Tables: {table_count}")
        
        # Look for placeholder patterns
        import re
        placeholder_patterns = re.findall(r'\{\{[^}]+\}\}', all_text)
        
        if placeholder_patterns:
            print(f"\n✓ Found {len(placeholder_patterns)} placeholder patterns:")
            # Remove duplicates and sort
            unique_placeholders = sorted(set(placeholder_patterns))
            for placeholder in unique_placeholders:
                print(f"  - {placeholder}")
        else:
            print("\n✗ No placeholder patterns ({{ }}) found")
        
        # Check for specific placeholders we expect
        expected_placeholders = [
            '{{TRIP_NUMBER}}', '{{PATIENT_NAME}}', '{{DEPARTURE_DATE}}',
            '{{ORIGIN_AIRPORT}}', '{{DESTINATION_AIRPORT}}', '{{AIRCRAFT_TAIL}}'
        ]
        
        found_expected = []
        for placeholder in expected_placeholders:
            if placeholder in all_text:
                found_expected.append(placeholder)
        
        if found_expected:
            print(f"\n✓ Found expected placeholders: {found_expected}")
        
        missing_expected = [p for p in expected_placeholders if p not in found_expected]
        if missing_expected:
            print(f"\n⚠️  Missing expected placeholders: {missing_expected}")
        
        # Show a sample of the content
        print(f"\nSample content (first 300 chars):")
        print(f"{all_text[:300]}...")
        
        # Check if this looks like a properly templated document
        has_placeholders = len(placeholder_patterns) > 0
        has_expected = len(found_expected) > 0
        
        if has_placeholders and has_expected:
            print(f"\n✅ RESULT: Document appears properly templated and ready for data population")
        elif has_placeholders:
            print(f"\n⚠️  RESULT: Document has placeholders but missing some expected ones")
        else:
            print(f"\n❌ RESULT: Document does not appear to be templated")
        
        return has_placeholders and has_expected
        
    except Exception as e:
        print(f"Error reading templated document: {e}")
        return False

if __name__ == "__main__":
    check_templated_document()