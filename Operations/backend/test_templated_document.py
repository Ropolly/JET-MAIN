#!/usr/bin/env python
"""
Test the templated CustomerItin document with actual trip data
"""

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

from api.models import Trip
from utils.docgen.trip_document_generator import TripDocumentGenerator

try:
    from docx import Document as DocxDocument
    print("python-docx is available")
except ImportError:
    print("python-docx not available")
    exit(1)

def test_templated_document():
    """Test document generation with the templated CustomerItin"""
    
    # Get first available trip
    trip = Trip.objects.first()
    if not trip:
        print("No trips available for testing")
        return
    
    print(f"Testing templated document generation for Trip {trip.trip_number} (ID: {trip.id})")
    
    # Create a modified generator that uses the templated directory
    generator = TripDocumentGenerator(str(trip.id))
    
    # Temporarily modify the template path to use the templated version
    original_templates_dir = generator.base_generator.templates_dir
    templated_dir = original_templates_dir / "templated"
    
    print(f"Using templated directory: {templated_dir}")
    
    # Check if templated CustomerItin exists
    templated_customer_itin = templated_dir / "CustomerItin.docx"
    if not templated_customer_itin.exists():
        print("Templated CustomerItin.docx not found")
        return
    
    try:
        # Get trip data
        trip_data = generator._prepare_trip_data()
        patient_data = generator._prepare_patient_data() if trip.patient else None
        
        # Create placeholder mappings
        placeholders = generator._create_placeholder_mappings(trip_data, None, patient_data)
        
        # Create a custom mapping that matches the template's lowercase placeholders
        custom_mappings = {
            '{{patient_name}}': placeholders.get('{{PATIENT_NAME}}', 'N/A'),
            '{{trip_date}}': placeholders.get('{{DEPARTURE_DATE}}', 'N/A'),
            '{{trip_number}}': placeholders.get('{{TRIP_NUMBER}}', 'N/A'),
            '{{trip_itinerary}}': f"{placeholders.get('{{ORIGIN_AIRPORT}}', 'N/A')} → {placeholders.get('{{DESTINATION_AIRPORT}}', 'N/A')}"
        }
        
        print(f"\nCustom mappings for template:")
        for placeholder, value in custom_mappings.items():
            print(f"  {placeholder}: {value}")
        
        # Generate document using templated version
        output_filename = f"TEMPLATED-{trip.trip_number}-customer_itinerary-test.docx"
        output_path = Path("documents") / output_filename
        
        # Use the template filling method
        filled_doc_path = generator._fill_template_with_data(
            str(templated_customer_itin),
            str(output_path),
            custom_mappings
        )
        
        print(f"\n✓ Generated templated document: {output_filename}")
        
        # Verify the generated document has data
        if Path(filled_doc_path).exists():
            doc = DocxDocument(filled_doc_path)
            all_text = ""
            
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
            
            print(f"\nGenerated document content verification:")
            
            # Check if our data made it in
            test_values = {
                trip.trip_number: "Trip number",
                placeholders.get('{{PATIENT_NAME}}', ''): "Patient name",
                placeholders.get('{{DEPARTURE_DATE}}', ''): "Departure date"
            }
            
            found_data = False
            for value, description in test_values.items():
                if value and value != 'N/A' and value in all_text:
                    print(f"  ✓ {description} found: {value}")
                    found_data = True
                elif value:
                    print(f"  ✗ {description} not found: {value}")
            
            # Check for unreplaced placeholders
            unreplaced = []
            for placeholder in custom_mappings.keys():
                if placeholder in all_text:
                    unreplaced.append(placeholder)
            
            if unreplaced:
                print(f"  ⚠️  Unreplaced placeholders: {unreplaced}")
            else:
                print(f"  ✓ All placeholders were replaced")
            
            if found_data and not unreplaced:
                print(f"\n🎉 SUCCESS: Templated document populated with trip data!")
                return True
            else:
                print(f"\n❌ Issues found with template data population")
                return False
        else:
            print(f"Generated document not found at: {filled_doc_path}")
            return False
            
    except Exception as e:
        print(f"Error testing templated document: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_templated_document()