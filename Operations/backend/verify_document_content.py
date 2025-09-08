#!/usr/bin/env python
"""
Script to verify document content has trip data filled in
"""

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

try:
    from docx import Document as DocxDocument
    print("python-docx is available")
except ImportError:
    print("python-docx not available")
    exit(1)

def verify_document_content():
    """Verify that generated documents contain trip data"""
    documents_dir = Path("documents")
    
    if not documents_dir.exists():
        print("Documents directory not found")
        return
    
    # Find the most recent documents
    docx_files = list(documents_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found")
        return
    
    # Check the most recent customer itinerary
    customer_itin_files = [f for f in docx_files if 'customer_itinerary' in f.name]
    if customer_itin_files:
        latest_itin = max(customer_itin_files, key=lambda f: f.stat().st_mtime)
        print(f"\nChecking: {latest_itin.name}")
        
        try:
            doc = DocxDocument(latest_itin)
            all_text = ""
            
            # Extract all text from document
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + "\n"
            
            # Check for trip data placeholders and actual data
            test_data = {
                '00002': 'Trip number found',
                'N36LJ': 'Aircraft tail number found',
                'Jack Al-Hussaini': 'Primary pilot name found',
                '2025-08-30': 'Departure date found',
                'JET ICU Medical Transport': 'Company name found'
            }
            
            print("Content verification:")
            for data, description in test_data.items():
                if data in all_text:
                    print(f"✓ {description}: {data}")
                else:
                    print(f"✗ {description}: NOT FOUND")
            
            # Check for unreplaced placeholders
            placeholders = ['{{TRIP_NUMBER}}', '{{AIRCRAFT_TAIL}}', '{{PRIMARY_PILOT}}']
            unreplaced = [p for p in placeholders if p in all_text]
            if unreplaced:
                print(f"\n⚠️  Unreplaced placeholders found: {unreplaced}")
            else:
                print(f"\n✓ No unreplaced placeholders detected")
                
        except Exception as e:
            print(f"Error reading document: {e}")
    
    # Check handling request as well
    handling_files = [f for f in docx_files if 'handling_request' in f.name]
    if handling_files:
        latest_handling = max(handling_files, key=lambda f: f.stat().st_mtime)
        print(f"\nChecking: {latest_handling.name}")
        
        try:
            doc = DocxDocument(latest_handling)
            all_text = ""
            
            # Extract all text from document
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + "\n"
            
            # Quick check for data
            if '00002' in all_text:
                print("✓ Trip number found in handling request")
            else:
                print("✗ Trip number not found in handling request")
                
        except Exception as e:
            print(f"Error reading handling request: {e}")

if __name__ == "__main__":
    verify_document_content()