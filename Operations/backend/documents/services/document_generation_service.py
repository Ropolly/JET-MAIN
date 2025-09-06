#!/usr/bin/env python
"""
Document Generation Service for JET Aviation Operations

Moved from utils/docgen/docgen.py to documents/services/document_generation_service.py
This service handles the generation of aviation-related documents including:
- General Declaration (GenDec)
- Handling Requests
- Customer Itineraries
- Internal Itineraries
- Quote Forms
"""

import os
from pathlib import Path
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, List
from django.utils import timezone
from django.conf import settings
from operations.models import TripLine, Trip
from aircraft.models import Aircraft
from contacts.models import Contact
from operations.models import Passenger, Patient
from ..models import Document, DocumentTemplate
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not installed. Install with: pip install python-docx")
    DocxDocument = None
    DOCX_AVAILABLE = False


class DocumentGenerationService:
    """Main document generation service for aviation documents."""
    
    def __init__(self):
        """Initialize the document generator with template and output paths."""
        self.base_dir = Path(__file__).parent.parent
        self.templates_dir = self.base_dir / "templates"
        self.outputs_dir = settings.MEDIA_ROOT / "generated_documents"
        
        # Ensure outputs directory exists
        self.outputs_dir.mkdir(parents=True, exist_ok=True)
        
        # Check if templates directory exists
        if not self.templates_dir.exists():
            self.templates_dir.mkdir(parents=True, exist_ok=True)
            logger.warning(f"Created templates directory: {self.templates_dir}")
    
    def generate_general_declaration(self, trip_id: str, output_filename: Optional[str] = None) -> Optional[Document]:
        """
        Generate a General Declaration (GenDec) document for a trip.
        
        Args:
            trip_id: UUID of the trip
            output_filename: Optional custom filename
            
        Returns:
            Document instance if successful, None otherwise
        """
        try:
            trip = Trip.objects.get(id=trip_id)
            
            if not DOCX_AVAILABLE:
                logger.error("python-docx not available for document generation")
                return None
            
            # Create new document
            doc = DocxDocument()
            
            # Add title
            title = doc.add_heading('GENERAL DECLARATION', 0)
            title.alignment = 1  # Center alignment
            
            # Add trip information
            doc.add_heading('Flight Information', level=1)
            
            # Basic trip details
            trip_table = doc.add_table(rows=0, cols=2)
            trip_table.style = 'Table Grid'
            
            self._add_table_row(trip_table, 'Trip Number:', trip.trip_number)
            self._add_table_row(trip_table, 'Trip Type:', trip.get_type_display())
            self._add_table_row(trip_table, 'Aircraft:', str(trip.aircraft) if trip.aircraft else 'TBD')
            
            if trip.estimated_departure_time:
                self._add_table_row(trip_table, 'Departure Time:', 
                                  trip.estimated_departure_time.strftime('%Y-%m-%d %H:%M UTC'))
            
            # Add trip lines (flight legs)
            if trip.trip_lines.exists():
                doc.add_heading('Flight Legs', level=1)
                
                legs_table = doc.add_table(rows=1, cols=5)
                legs_table.style = 'Table Grid'
                
                # Header row
                header_cells = legs_table.rows[0].cells
                header_cells[0].text = 'Origin'
                header_cells[1].text = 'Destination'
                header_cells[2].text = 'Departure'
                header_cells[3].text = 'Arrival'
                header_cells[4].text = 'Distance'
                
                for trip_line in trip.trip_lines.all():
                    row_cells = legs_table.add_row().cells
                    row_cells[0].text = str(trip_line.origin_airport)
                    row_cells[1].text = str(trip_line.destination_airport)
                    row_cells[2].text = trip_line.departure_time_local.strftime('%Y-%m-%d %H:%M')
                    row_cells[3].text = trip_line.arrival_time_local.strftime('%Y-%m-%d %H:%M')
                    row_cells[4].text = f"{trip_line.distance} nm"
            
            # Add passengers if any
            if trip.passengers.exists():
                doc.add_heading('Passengers', level=1)
                
                pax_table = doc.add_table(rows=1, cols=4)
                pax_table.style = 'Table Grid'
                
                # Header row
                header_cells = pax_table.rows[0].cells
                header_cells[0].text = 'Name'
                header_cells[1].text = 'Nationality'
                header_cells[2].text = 'Passport'
                header_cells[3].text = 'Date of Birth'
                
                for passenger in trip.passengers.all():
                    row_cells = pax_table.add_row().cells
                    row_cells[0].text = str(passenger.info)
                    row_cells[1].text = passenger.nationality or 'N/A'
                    row_cells[2].text = passenger.passport_number or 'N/A'
                    row_cells[3].text = passenger.date_of_birth.strftime('%Y-%m-%d') if passenger.date_of_birth else 'N/A'
            
            # Add patient information if medical trip
            if trip.patient:
                doc.add_heading('Patient Information', level=1)
                
                patient_table = doc.add_table(rows=0, cols=2)
                patient_table.style = 'Table Grid'
                
                self._add_table_row(patient_table, 'Patient Name:', str(trip.patient.info))
                self._add_table_row(patient_table, 'Date of Birth:', 
                                  trip.patient.date_of_birth.strftime('%Y-%m-%d'))
                self._add_table_row(patient_table, 'Nationality:', trip.patient.nationality)
                self._add_table_row(patient_table, 'Passport:', trip.patient.passport_number)
                
                if trip.patient.special_instructions:
                    doc.add_heading('Special Instructions', level=2)
                    doc.add_paragraph(trip.patient.special_instructions)
            
            # Save document
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"GenDec_{trip.trip_number}_{timestamp}.docx"
            
            output_path = self.outputs_dir / output_filename
            doc.save(str(output_path))
            
            # Create Document record
            with open(output_path, 'rb') as f:
                content = f.read()
            
            document = Document.objects.create(
                filename=output_filename,
                content=content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                file_size=len(content),
                document_category='itinerary'
            )
            
            logger.info(f"Generated General Declaration: {output_filename}")
            return document
            
        except Trip.DoesNotExist:
            logger.error(f"Trip not found: {trip_id}")
            return None
        except Exception as e:
            logger.error(f"Error generating General Declaration: {str(e)}")
            return None
    
    def generate_quote_document(self, quote_id: str, output_filename: Optional[str] = None) -> Optional[Document]:
        """
        Generate a quote document.
        
        Args:
            quote_id: UUID of the quote
            output_filename: Optional custom filename
            
        Returns:
            Document instance if successful, None otherwise
        """
        try:
            from operations.models import Quote
            quote = Quote.objects.get(id=quote_id)
            
            if not DOCX_AVAILABLE:
                logger.error("python-docx not available for document generation")
                return None
            
            # Create new document
            doc = DocxDocument()
            
            # Add title
            title = doc.add_heading('FLIGHT QUOTE', 0)
            title.alignment = 1  # Center alignment
            
            # Add quote information
            doc.add_heading('Quote Details', level=1)
            
            quote_table = doc.add_table(rows=0, cols=2)
            quote_table.style = 'Table Grid'
            
            self._add_table_row(quote_table, 'Quote Amount:', f"${quote.quoted_amount:,.2f}")
            self._add_table_row(quote_table, 'Contact:', str(quote.contact))
            self._add_table_row(quote_table, 'Aircraft Type:', quote.get_aircraft_type_display())
            self._add_table_row(quote_table, 'Medical Team:', quote.get_medical_team_display())
            self._add_table_row(quote_table, 'Pickup Airport:', str(quote.pickup_airport))
            self._add_table_row(quote_table, 'Dropoff Airport:', str(quote.dropoff_airport))
            self._add_table_row(quote_table, 'Estimated Flight Time:', str(quote.estimated_flight_time))
            self._add_table_row(quote_table, 'Number of Stops:', str(quote.number_of_stops))
            self._add_table_row(quote_table, 'Includes Ground Transport:', 'Yes' if quote.includes_grounds else 'No')
            self._add_table_row(quote_table, 'Status:', quote.get_status_display())
            self._add_table_row(quote_table, 'Payment Status:', quote.get_payment_status_display())
            
            # Add cruise information if available
            if quote.cruise_line:
                doc.add_heading('Cruise Information', level=1)
                
                cruise_table = doc.add_table(rows=0, cols=2)
                cruise_table.style = 'Table Grid'
                
                self._add_table_row(cruise_table, 'Cruise Line:', quote.cruise_line)
                self._add_table_row(cruise_table, 'Ship:', quote.cruise_ship or 'N/A')
                if quote.cruise_doctor_first_name and quote.cruise_doctor_last_name:
                    doctor_name = f"{quote.cruise_doctor_first_name} {quote.cruise_doctor_last_name}"
                    self._add_table_row(cruise_table, 'Ship Doctor:', doctor_name)
            
            # Add patient information if available
            if quote.patient:
                doc.add_heading('Patient Information', level=1)
                
                patient_table = doc.add_table(rows=0, cols=2)
                patient_table.style = 'Table Grid'
                
                self._add_table_row(patient_table, 'Patient:', str(quote.patient.info))
                self._add_table_row(patient_table, 'Bed at Origin:', 'Yes' if quote.patient.bed_at_origin else 'No')
                self._add_table_row(patient_table, 'Bed at Destination:', 'Yes' if quote.patient.bed_at_destination else 'No')
                
                if quote.patient.special_instructions:
                    doc.add_heading('Special Instructions', level=2)
                    doc.add_paragraph(quote.patient.special_instructions)
            
            # Add footer with generation date
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            footer.alignment = 1  # Center alignment
            
            # Save document
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"Quote_{quote.id}_{timestamp}.docx"
            
            output_path = self.outputs_dir / output_filename
            doc.save(str(output_path))
            
            # Create Document record
            with open(output_path, 'rb') as f:
                content = f.read()
            
            document = Document.objects.create(
                filename=output_filename,
                content=content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                file_size=len(content),
                document_category='quote'
            )
            
            logger.info(f"Generated quote document: {output_filename}")
            return document
            
        except Exception as e:
            logger.error(f"Error generating quote document: {str(e)}")
            return None
    
    def generate_itinerary(self, trip_id: str, customer_facing: bool = True, 
                          output_filename: Optional[str] = None) -> Optional[Document]:
        """
        Generate an itinerary document for a trip.
        
        Args:
            trip_id: UUID of the trip
            customer_facing: Whether this is for customer or internal use
            output_filename: Optional custom filename
            
        Returns:
            Document instance if successful, None otherwise
        """
        try:
            trip = Trip.objects.get(id=trip_id)
            
            if not DOCX_AVAILABLE:
                logger.error("python-docx not available for document generation")
                return None
            
            # Create new document
            doc = DocxDocument()
            
            # Add title
            title_text = 'FLIGHT ITINERARY' if customer_facing else 'INTERNAL FLIGHT ITINERARY'
            title = doc.add_heading(title_text, 0)
            title.alignment = 1  # Center alignment
            
            # Add trip information
            doc.add_heading('Trip Information', level=1)
            
            trip_table = doc.add_table(rows=0, cols=2)
            trip_table.style = 'Table Grid'
            
            self._add_table_row(trip_table, 'Trip Number:', trip.trip_number)
            self._add_table_row(trip_table, 'Trip Type:', trip.get_type_display())
            
            if not customer_facing or trip.aircraft:
                self._add_table_row(trip_table, 'Aircraft:', str(trip.aircraft) if trip.aircraft else 'TBD')
            
            if trip.estimated_departure_time:
                self._add_table_row(trip_table, 'Departure Time:', 
                                  trip.estimated_departure_time.strftime('%Y-%m-%d %H:%M UTC'))
            
            # Add detailed flight schedule
            if trip.trip_lines.exists():
                doc.add_heading('Flight Schedule', level=1)
                
                for i, trip_line in enumerate(trip.trip_lines.all(), 1):
                    doc.add_heading(f'Leg {i}: {trip_line.origin_airport} → {trip_line.destination_airport}', level=2)
                    
                    leg_table = doc.add_table(rows=0, cols=2)
                    leg_table.style = 'Table Grid'
                    
                    self._add_table_row(leg_table, 'Origin:', f"{trip_line.origin_airport} ({trip_line.origin_airport.city})")
                    self._add_table_row(leg_table, 'Destination:', f"{trip_line.destination_airport} ({trip_line.destination_airport.city})")
                    self._add_table_row(leg_table, 'Departure (Local):', trip_line.departure_time_local.strftime('%Y-%m-%d %H:%M'))
                    self._add_table_row(leg_table, 'Arrival (Local):', trip_line.arrival_time_local.strftime('%Y-%m-%d %H:%M'))
                    self._add_table_row(leg_table, 'Flight Time:', str(trip_line.flight_time))
                    self._add_table_row(leg_table, 'Distance:', f"{trip_line.distance} nm")
                    
                    if not customer_facing:
                        self._add_table_row(leg_table, 'Departure (UTC):', trip_line.departure_time_utc.strftime('%Y-%m-%d %H:%M'))
                        self._add_table_row(leg_table, 'Arrival (UTC):', trip_line.arrival_time_utc.strftime('%Y-%m-%d %H:%M'))
                        
                        if trip_line.departure_fbo:
                            self._add_table_row(leg_table, 'Departure FBO:', str(trip_line.departure_fbo))
                        if trip_line.arrival_fbo:
                            self._add_table_row(leg_table, 'Arrival FBO:', str(trip_line.arrival_fbo))
                        
                        if trip_line.crew_line:
                            self._add_table_row(leg_table, 'Crew:', str(trip_line.crew_line))
            
            # Add notes if any
            if trip.notes:
                doc.add_heading('Notes', level=1)
                doc.add_paragraph(trip.notes)
            
            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            footer.alignment = 1  # Center alignment
            
            # Save document
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                doc_type = 'Customer' if customer_facing else 'Internal'
                output_filename = f"{doc_type}_Itinerary_{trip.trip_number}_{timestamp}.docx"
            
            output_path = self.outputs_dir / output_filename
            doc.save(str(output_path))
            
            # Create Document record
            with open(output_path, 'rb') as f:
                content = f.read()
            
            document = Document.objects.create(
                filename=output_filename,
                content=content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                file_size=len(content),
                document_category='itinerary'
            )
            
            logger.info(f"Generated itinerary: {output_filename}")
            return document
            
        except Trip.DoesNotExist:
            logger.error(f"Trip not found: {trip_id}")
            return None
        except Exception as e:
            logger.error(f"Error generating itinerary: {str(e)}")
            return None
    
    def _add_table_row(self, table, label: str, value: str):
        """Helper method to add a row to a table."""
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value
    
    def list_available_templates(self) -> List[str]:
        """List all available document templates."""
        if not self.templates_dir.exists():
            return []
        
        templates = []
        for file_path in self.templates_dir.glob('*.docx'):
            templates.append(file_path.name)
        
        return templates
    
    def cleanup_old_documents(self, days_old: int = 30):
        """Clean up old generated documents."""
        cutoff_date = datetime.now() - timedelta(days=days_old)
        
        deleted_count = 0
        for file_path in self.outputs_dir.glob('*.docx'):
            if file_path.stat().st_mtime < cutoff_date.timestamp():
                try:
                    file_path.unlink()
                    deleted_count += 1
                except OSError as e:
                    logger.error(f"Error deleting old document {file_path}: {e}")
        
        logger.info(f"Cleaned up {deleted_count} old documents")
        return deleted_count


# Convenience functions for backward compatibility
def generate_general_declaration(trip_id: str, output_filename: Optional[str] = None) -> Optional[Document]:
    """Generate a General Declaration document."""
    service = DocumentGenerationService()
    return service.generate_general_declaration(trip_id, output_filename)


def generate_quote_document(quote_id: str, output_filename: Optional[str] = None) -> Optional[Document]:
    """Generate a quote document."""
    service = DocumentGenerationService()
    return service.generate_quote_document(quote_id, output_filename)


def generate_itinerary(trip_id: str, customer_facing: bool = True, 
                      output_filename: Optional[str] = None) -> Optional[Document]:
    """Generate an itinerary document."""
    service = DocumentGenerationService()
    return service.generate_itinerary(trip_id, customer_facing, output_filename)
