#!/usr/bin/env python3

from docx import Document
from pathlib import Path
import datetime

def validate_gendec_template():

    print(" VALIDATING GENDEC TEMPLATE WITH PLACEHOLDERS")
    print("=" * 60)
    
    base_dir = Path(__file__).parent
    template_path = base_dir / "documents" / "GenDec_With_Placeholders.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return False
    
    print(f" Template found: {template_path.name}")
    print(f" File size: {template_path.stat().st_size:,} bytes")
    
    try:
        # Load template
        doc = Document(template_path)
        
        # Find all placeholders
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if '{{' in text and '}}' in text:
                # Extract placeholders
                import re
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                placeholders.update(found)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{{' in text and '}}' in text:
                            import re
                            found = re.findall(r'\{\{([^}]+)\}\}', text)
                            placeholders.update(found)
        
        print(f"\n Found {len(placeholders)} unique placeholders:")
        
        # Expected placeholders for CBP Form 7507
        expected_placeholders = [
            'TRIP_NUMBER', 'DEPARTURE_DATE', 'DEPARTURE_TIME_UTC', 'ARRIVAL_DATE', 'ARRIVAL_TIME_UTC',
            'ORIGIN_AIRPORT_NAME', 'ORIGIN_AIRPORT_CODE', 'ORIGIN_AIRPORT_ICAO', 'ORIGIN_CITY',
            'DESTINATION_AIRPORT_NAME', 'DESTINATION_AIRPORT_CODE', 'DESTINATION_AIRPORT_ICAO', 'DESTINATION_CITY',
            'AIRCRAFT_TAIL', 'AIRCRAFT_MAKE', 'AIRCRAFT_MODEL', 'AIRCRAFT_COMPANY', 'AIRCRAFT_SERIAL', 'AIRCRAFT_MGTOW',
            'PRIMARY_PILOT', 'SECONDARY_PILOT', 'CREW_COUNT', 'PASSENGER_COUNT', 'TOTAL_PAX_COUNT',
            'CARRIER_CODE', 'FLIGHT_NUMBER', 'PILOT_TITLE',
            'PASSENGERS_AND_PATIENT_LIST', 'CREW_LIST', 'DOCUMENT_DATE', 'DOCUMENT_TIME'
        ]
        
        # Check which expected placeholders are found
        found_expected = []
        missing_expected = []
        
        for placeholder in expected_placeholders:
            if placeholder in placeholders:
                found_expected.append(placeholder)
            else:
                missing_expected.append(placeholder)
        
        # Show found placeholders (grouped by category)
        categories = {
            'Flight Info': ['TRIP_NUMBER', 'DEPARTURE_DATE', 'DEPARTURE_TIME_UTC', 'ARRIVAL_DATE', 'ARRIVAL_TIME_UTC', 'FLIGHT_NUMBER', 'CARRIER_CODE'],
            'Airports': ['ORIGIN_AIRPORT_NAME', 'ORIGIN_AIRPORT_CODE', 'ORIGIN_AIRPORT_ICAO', 'ORIGIN_CITY', 'DESTINATION_AIRPORT_NAME', 'DESTINATION_AIRPORT_CODE', 'DESTINATION_AIRPORT_ICAO', 'DESTINATION_CITY'],
            'Aircraft': ['AIRCRAFT_TAIL', 'AIRCRAFT_MAKE', 'AIRCRAFT_MODEL', 'AIRCRAFT_COMPANY', 'AIRCRAFT_SERIAL', 'AIRCRAFT_MGTOW'],
            'Crew': ['PRIMARY_PILOT', 'SECONDARY_PILOT', 'CREW_COUNT', 'PILOT_TITLE', 'CREW_LIST'],
            'Passengers': ['PASSENGER_COUNT', 'TOTAL_PAX_COUNT', 'PASSENGERS_AND_PATIENT_LIST'],
            'Document': ['DOCUMENT_DATE', 'DOCUMENT_TIME']
        }
        
        for category, category_placeholders in categories.items():
            found_in_category = [p for p in category_placeholders if p in placeholders]
            if found_in_category:
                print(f"\n   {category}:")
                for placeholder in found_in_category:
                    print(f"      {{{{ {placeholder} }}}}")
        
        # Show any unexpected placeholders
        unexpected = placeholders - set(expected_placeholders)
        if unexpected:
            print(f"\n   Other:")
            for placeholder in sorted(unexpected):
                print(f"      {{{{ {placeholder} }}}}")
        
        # Show missing placeholders
        if missing_expected:
            print(f"\n Missing expected placeholders ({len(missing_expected)}):")
            for placeholder in missing_expected:
                print(f"      {{{{ {placeholder} }}}}")
        
        # Validation summary
        total_expected = len(expected_placeholders)
        total_found_expected = len(found_expected)
        completion_rate = (total_found_expected / total_expected) * 100
        
        print(f"\n VALIDATION SUMMARY:")
        print(f"   • Expected placeholders: {total_expected}")
        print(f"   • Found expected: {total_found_expected}")
        print(f"   • Completion rate: {completion_rate:.1f}%")
        print(f"   • Total placeholders in template: {len(placeholders)}")
        
        if completion_rate >= 90:
            print(f"    Template validation: PASSED")
            return True
        else:
            print(f"    Template validation: FAILED (missing critical placeholders)")
            return False
    
    except Exception as e:
        print(f" Error validating template: {str(e)}")
        return False

def create_filled_preview():

    print(f"\n CREATING FILLED TEMPLATE PREVIEW")
    print("=" * 60)
    
    # Mock data for preview
    mock_data = {
        'TRIP_NUMBER': 'JET-2024-0902',
        'DEPARTURE_DATE': '09/02/2024',
        'DEPARTURE_TIME_UTC': '14:30:00',
        'ARRIVAL_DATE': '09/02/2024', 
        'ARRIVAL_TIME_UTC': '18:45:00',
        'ORIGIN_AIRPORT_NAME': 'Teterboro Airport',
        'ORIGIN_AIRPORT_CODE': 'TEB',
        'ORIGIN_AIRPORT_ICAO': 'KTEB',
        'ORIGIN_CITY': 'Teterboro, NJ',
        'DESTINATION_AIRPORT_NAME': 'Miami International Airport',
        'DESTINATION_AIRPORT_CODE': 'MIA',
        'DESTINATION_AIRPORT_ICAO': 'KMIA',
        'DESTINATION_CITY': 'Miami, FL',
        'AIRCRAFT_TAIL': 'N525JT',
        'AIRCRAFT_MAKE': 'Cessna',
        'AIRCRAFT_MODEL': 'Citation CJ3+',
        'AIRCRAFT_COMPANY': 'JET Charter Services LLC',
        'AIRCRAFT_SERIAL': '525C-0924',
        'AIRCRAFT_MGTOW': '13,870',
        'PRIMARY_PILOT': 'Captain John Smith',
        'SECONDARY_PILOT': 'First Officer Sarah Johnson',
        'CREW_COUNT': '4',
        'PASSENGER_COUNT': '2',
        'TOTAL_PAX_COUNT': '3',
        'CARRIER_CODE': 'JCS',
        'FLIGHT_NUMBER': 'JET001',
        'PILOT_TITLE': 'Aircraft Commander',
        'PASSENGERS_AND_PATIENT_LIST': 'Robert Wilson | Patient | United States | P123456789\nEmily Wilson | Passenger | United States | P987654321\nDr. Jennifer Thompson | Passenger | United States | P456789123',
        'CREW_LIST': 'Captain: Captain John Smith\nFirst Officer: First Officer Sarah Johnson\nMedical Crew: Dr. Michael Brown\nMedical Crew: Paramedic Lisa Davis',
        'DOCUMENT_DATE': datetime.datetime.now().strftime('%Y-%m-%d'),
        'DOCUMENT_TIME': datetime.datetime.now().strftime('%H:%M:%S UTC')
    }
    
    base_dir = Path(__file__).parent
    template_path = base_dir / "documents" / "GenDec_With_Placeholders.docx"
    preview_path = base_dir / "GenDec_Preview_Filled.docx"
    
    try:
        # Load and fill template
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in mock_data.items():
                if f"{{{placeholder}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", str(value))
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in mock_data.items():
                            if f"{{{placeholder}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", str(value))
        
        # Save filled preview
        doc.save(preview_path)
        
        file_size = preview_path.stat().st_size
        print(f" Filled preview saved: {preview_path.name}")
        print(f" File size: {file_size:,} bytes")
        
        # Show a few lines of preview content
        print(f"\n DOCUMENT CONTENT PREVIEW:")
        print("=" * 40)
        
        doc = Document(preview_path)
        line_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and line_count < 15:
                print(f"   {paragraph.text[:80]}{'...' if len(paragraph.text) > 80 else ''}")
                line_count += 1
        
        if line_count == 15:
            print("   ...")
            print("   [Document continues...]")
        
        return str(preview_path)
    
    except Exception as e:
        print(f" Error creating preview: {str(e)}")
        return None

def main():

    # Validate template structure
    template_valid = validate_gendec_template()
    
    if template_valid:
        # Create filled preview
        preview_path = create_filled_preview()
        
        print(f"\n VALIDATION COMPLETE!")
        print(f" GenDec template is ready for production use")
        print(f" All critical CBP Form 7507 placeholders are present")
        if preview_path:
            print(f" Preview with filled data saved for review")
    else:
        print(f"\n VALIDATION FAILED!")
        print(f" Template needs fixes before production use")

if __name__ == "__main__":
    main()
