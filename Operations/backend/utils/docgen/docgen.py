#!/usr/bin/env python

import os
import sys
import shutil
import django
from pathlib import Path
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, List

# Import PDF generator for GenDec
try:
    import fitz  # PyMuPDF for PDF form filling
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    fitz = None

# Setup Django environment for standalone script execution
sys.path.append(str(Path(__file__).parent.parent.parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

from django.utils import timezone
from django.http import HttpResponse
from api.models import TripLine, Trip, Aircraft, Contact, Passenger, Patient
import io

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Warning: python-docx not installed. Install with: pip install python-docx")
    Document = None

class DocumentGenerator:
    
    def __init__(self):
        self.base_dir = Path(__file__).parent
        self.templates_dir = self.base_dir / "documents"
        self.outputs_dir = self.base_dir / "outputs"
        
        # Ensure outputs directory exists
        self.outputs_dir.mkdir(exist_ok=True)
        
        # Check if templates directory exists
        if not self.templates_dir.exists():
            raise FileNotFoundError(f"Templates directory not found: {self.templates_dir}")
    
    def generate_general_declaration(self, trip_line_id: str) -> str:

        try:
            trip_line = TripLine.objects.select_related(
                'trip', 'trip__aircraft', 'trip__patient', 'trip__patient__info',
                'origin_airport', 'destination_airport', 'crew_line',
                'crew_line__primary_in_command', 'crew_line__secondary_in_command'
            ).prefetch_related(
                'trip__passengers', 'trip__passengers__info',
                'crew_line__medic_ids'
            ).get(id=trip_line_id)
        except TripLine.DoesNotExist:
            raise TripLine.DoesNotExist(f"Trip line with ID {trip_line_id} not found")
        
        # Prepare data for GenDec
        data = self._prepare_gendec_data(trip_line)
        
        # Use the original GenDec.docx template (now contains form fields with placeholders)
        template_path = self.templates_dir / "GenDec.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"GenDec template not found: {template_path}")
        
        # Create output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        trip_number = trip_line.trip.trip_number or "UNKNOWN"
        leg_suffix = f"{trip_line.origin_airport.iata_code}_{trip_line.destination_airport.iata_code}"
        output_filename = f"GenDec_{trip_number}_{leg_suffix}_{timestamp}.docx"
        output_path = self.outputs_dir / output_filename
        
        # Generate the document
        if Document is not None:
            try:
                return self._fill_docx_template(template_path, output_path, data)
            except Exception as e:
                print(f"Error using python-docx: {e}. Falling back to text replacement.")
                return self._fill_template_fallback(template_path, output_path, data)
        else:
            return self._fill_template_fallback(template_path, output_path, data)
    
    def generate_gendec_pdf(self, trip_line_id: str) -> dict:

        if not PDF_AVAILABLE:
            raise Exception("PyMuPDF not available - cannot generate PDF forms")
            
        try:
            trip_line = TripLine.objects.select_related(
                'trip', 'trip__aircraft', 'trip__patient', 'trip__patient__info',
                'origin_airport', 'destination_airport', 'crew_line',
                'crew_line__primary_in_command', 'crew_line__secondary_in_command'
            ).prefetch_related(
                'trip__passengers', 'trip__passengers__info',
                'crew_line__medic_ids'
            ).get(id=trip_line_id)
        except TripLine.DoesNotExist:
            raise TripLine.DoesNotExist(f"Trip line with ID {trip_line_id} not found")
        
        # Check for PDF template
        template_path = self.templates_dir / "GenDec.pdf"
        if not template_path.exists():
            raise FileNotFoundError(f"GenDec.pdf template not found: {template_path}")
        
        return self._fill_gendec_pdf(trip_line, template_path)

    def generate_handling_request(self, trip_line_id: str) -> str:

        try:
            trip_line = TripLine.objects.select_related(
                'trip', 'trip__aircraft', 'trip__patient', 'trip__patient__info',
                'origin_airport', 'destination_airport', 'crew_line',
                'crew_line__primary_in_command', 'crew_line__secondary_in_command'
            ).prefetch_related(
                'trip__passengers', 'trip__passengers__info',
                'crew_line__medic_ids'
            ).get(id=trip_line_id)
        except TripLine.DoesNotExist:
            raise TripLine.DoesNotExist(f"Trip line with ID {trip_line_id} not found")
        
        # Prepare data using same logic as GenDec
        data = self._prepare_gendec_data(trip_line)
        
        # Generate document
        template_path = self.templates_dir / "HandlingRequest.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Handling Request template not found: {template_path}")
        
        # Create output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        trip_number = trip_line.trip.trip_number or "UNKNOWN"
        leg_suffix = f"{trip_line.origin_airport.iata_code}_{trip_line.destination_airport.iata_code}"
        output_filename = f"HandlingRequest_{trip_number}_{leg_suffix}_{timestamp}.docx"
        output_path = self.outputs_dir / output_filename
        
        # Generate the document
        if Document is not None:
            try:
                return self._fill_docx_template(template_path, output_path, data)
            except Exception as e:
                print(f"Error using python-docx: {e}. Falling back to text replacement.")
                return self._fill_template_fallback(template_path, output_path, data)
        else:
            return self._fill_template_fallback(template_path, output_path, data)
    
    def generate_gendec_response(self, trip_line_id: str) -> HttpResponse:

        try:
            trip_line = TripLine.objects.select_related(
                'trip', 'trip__aircraft', 'trip__patient', 'trip__patient__info',
                'origin_airport', 'destination_airport', 'crew_line',
                'crew_line__primary_in_command', 'crew_line__secondary_in_command'
            ).prefetch_related(
                'trip__passengers', 'trip__passengers__info',
                'crew_line__medic_ids'
            ).get(id=trip_line_id)
        except TripLine.DoesNotExist:
            raise TripLine.DoesNotExist(f"Trip line with ID {trip_line_id} not found")
        
        # Prepare data
        data = self._prepare_gendec_data(trip_line)
        
        # Generate document in memory
        template_path = self.templates_dir / "GenDec.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"GenDec template not found: {template_path}")
        
        # Create filename for download
        trip_number = trip_line.trip.trip_number or "UNKNOWN"
        leg_suffix = f"{trip_line.origin_airport.iata_code}_{trip_line.destination_airport.iata_code}"
        filename = f"GenDec_{trip_number}_{leg_suffix}.docx"
        
        # Generate document in memory
        doc_content = self._generate_document_in_memory(template_path, data)
        
        # Create HTTP response
        response = HttpResponse(
            doc_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    def generate_handling_request_response(self, trip_line_id: str) -> HttpResponse:

        try:
            trip_line = TripLine.objects.select_related(
                'trip', 'trip__aircraft', 'trip__patient', 'trip__patient__info',
                'origin_airport', 'destination_airport', 'crew_line',
                'crew_line__primary_in_command', 'crew_line__secondary_in_command'
            ).prefetch_related(
                'trip__passengers', 'trip__passengers__info',
                'crew_line__medic_ids'
            ).get(id=trip_line_id)
        except TripLine.DoesNotExist:
            raise TripLine.DoesNotExist(f"Trip line with ID {trip_line_id} not found")
        
        # Prepare data using same logic as GenDec
        data = self._prepare_gendec_data(trip_line)
        
        # Generate document in memory
        template_path = self.templates_dir / "HandlingRequest.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Handling Request template not found: {template_path}")
        
        # Create filename for download
        trip_number = trip_line.trip.trip_number or "UNKNOWN"
        leg_suffix = f"{trip_line.origin_airport.iata_code}_{trip_line.destination_airport.iata_code}"
        filename = f"HandlingRequest_{trip_number}_{leg_suffix}.docx"
        
        # Generate document in memory
        doc_content = self._generate_document_in_memory(template_path, data)
        
        # Create HTTP response
        response = HttpResponse(
            doc_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    def _generate_document_in_memory(self, template_path: Path, data: Dict[str, Any]) -> bytes:

        if Document is not None:
            try:
                # Load template
                doc = Document(template_path)
                
                # Replace placeholders in paragraphs
                for paragraph in doc.paragraphs:
                    self._replace_placeholders_in_text(paragraph, data)
                
                # Replace placeholders in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_placeholders_in_text(paragraph, data)
                
                # Save to memory buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()
            except Exception as e:
                print(f"Error using python-docx: {e}. Using fallback method.")
                return self._generate_document_fallback(template_path)
        else:
            return self._generate_document_fallback(template_path)
    
    def _generate_document_fallback(self, template_path: Path) -> bytes:

        with open(template_path, 'rb') as f:
            return f.read()
    
    def _prepare_gendec_data(self, trip_line: TripLine) -> Dict[str, Any]:

        trip = trip_line.trip
        aircraft = trip.aircraft
        origin = trip_line.origin_airport
        destination = trip_line.destination_airport
        crew_line = trip_line.crew_line
        
        # Format dates and times
        departure_local = trip_line.departure_time_local.strftime("%Y-%m-%d %H:%M") if trip_line.departure_time_local else "TBD"
        departure_utc = trip_line.departure_time_utc.strftime("%Y-%m-%d %H:%M UTC") if trip_line.departure_time_utc else "TBD"
        arrival_local = trip_line.arrival_time_local.strftime("%Y-%m-%d %H:%M") if trip_line.arrival_time_local else "TBD"
        arrival_utc = trip_line.arrival_time_utc.strftime("%Y-%m-%d %H:%M UTC") if trip_line.arrival_time_utc else "TBD"
        
        # Aircraft information
        aircraft_info = {
            'tail_number': aircraft.tail_number if aircraft else "N/A",
            'make': aircraft.make if aircraft else "N/A",
            'model': aircraft.model if aircraft else "N/A",
            'company': aircraft.company if aircraft else "N/A",
            'serial_number': aircraft.serial_number if aircraft else "N/A",
            'mgtow': str(aircraft.mgtow) if aircraft else "N/A"
        }
        
        # Crew information
        crew_info = {
            'primary_pilot': self._format_contact(crew_line.primary_in_command) if crew_line and crew_line.primary_in_command else "N/A",
            'secondary_pilot': self._format_contact(crew_line.secondary_in_command) if crew_line and crew_line.secondary_in_command else "N/A",
            'medics': []
        }
        
        if crew_line:
            for medic in crew_line.medic_ids.all():
                crew_info['medics'].append(self._format_contact(medic))
        
        # Passenger and patient information (combined for PAX legs)
        passengers_and_patient = []
        total_pax_count = 0
        
        if trip_line.passenger_leg:
            # Add patient first if exists
            if trip.patient:
                patient_data = {
                    'name': self._format_contact(trip.patient.info) if trip.patient.info else "N/A",
                    'type': 'Patient',
                    'nationality': getattr(trip.patient, 'nationality', 'N/A'),
                    'passport_number': getattr(trip.patient, 'passport_number', 'N/A'),
                    'date_of_birth': trip.patient.date_of_birth.strftime("%Y-%m-%d") if trip.patient.date_of_birth else "N/A"
                }
                passengers_and_patient.append(patient_data)
                total_pax_count += 1
            
            # Add all passengers
            if trip.passengers.exists():
                for passenger in trip.passengers.all():
                    passenger_data = {
                        'name': self._format_contact(passenger.info),
                        'type': 'Passenger',
                        'nationality': passenger.nationality or "N/A",
                        'passport_number': passenger.passport_number or "N/A",
                        'date_of_birth': passenger.date_of_birth.strftime("%Y-%m-%d") if passenger.date_of_birth else "N/A"
                    }
                    passengers_and_patient.append(passenger_data)
                    total_pax_count += 1
        
        # Legacy separate fields for backward compatibility
        passengers_info = [p for p in passengers_and_patient if p.get('type') == 'Passenger']
        patient_info = next((p for p in passengers_and_patient if p.get('type') == 'Patient'), None)
        
        # Compile all data
        data = {
            # Flight information
            'trip_number': trip.trip_number or "N/A",
            'flight_type': "Passenger" if trip_line.passenger_leg else "Repositioning",
            'departure_date': departure_local,
            'departure_time_utc': departure_utc,
            'arrival_date': arrival_local,
            'arrival_time_utc': arrival_utc,
            'flight_time': str(trip_line.flight_time) if trip_line.flight_time else "N/A",
            'distance': f"{trip_line.distance} NM" if trip_line.distance else "N/A",
            
            # Airport information
            'origin_airport_name': origin.name if origin else "N/A",
            'origin_airport_code': origin.iata_code if origin else "N/A",
            'origin_airport_icao': origin.icao_code if origin else "N/A",
            'origin_city': origin.city if origin else "N/A",
            'origin_country': origin.country if origin else "N/A",
            
            'destination_airport_name': destination.name if destination else "N/A",
            'destination_airport_code': destination.iata_code if destination else "N/A",
            'destination_airport_icao': destination.icao_code if destination else "N/A",
            'destination_city': destination.city if destination else "N/A",
            'destination_country': destination.country if destination else "N/A",
            
            # Aircraft and crew
            'aircraft': aircraft_info,
            'crew': crew_info,
            'crew_count': self._calculate_crew_count(crew_line),
            
            # Passengers and patient (combined for PAX legs)
            'passengers': passengers_info,
            'passenger_count': len(passengers_info),
            'patient': patient_info,
            'passengers_and_patient': passengers_and_patient,
            'total_pax_count': total_pax_count,
            'is_pax_leg': trip_line.passenger_leg,
            
            # Document metadata
            'document_date': timezone.now().strftime("%Y-%m-%d"),
            'document_time': timezone.now().strftime("%H:%M:%S UTC")
        }
        
        return data
    
    def _format_contact(self, contact: Contact) -> str:

        if not contact:
            return "N/A"
        
        first_name = contact.first_name or ""
        last_name = contact.last_name or ""
        
        if first_name and last_name:
            return f"{first_name} {last_name}"
        elif first_name:
            return first_name
        elif last_name:
            return last_name
        else:
            return contact.email or "N/A"
    
    def _calculate_crew_count(self, crew_line) -> int:

        if not crew_line:
            return 0
            
        crew_count = 0
        
        # Count primary and secondary pilots
        if crew_line.primary_in_command:
            crew_count += 1
        if crew_line.secondary_in_command:
            crew_count += 1
            
        # Count medics if any
        if hasattr(crew_line, 'medic') and crew_line.medic.exists():
            crew_count += crew_line.medic.count()
            
        # Minimum crew count is 1 (at least a pilot)
        return max(crew_count, 1)
    
    def _fill_docx_template(self, template_path: Path, output_path: Path, data: Dict[str, Any]) -> str:

        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_text(paragraph, data)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_text(paragraph, data)
        
        # Save the document
        doc.save(output_path)
        return str(output_path)
    
    def _replace_placeholders_in_text(self, paragraph, data: Dict[str, Any]):

        # Create mapping of placeholders to actual values
        replacements = {
            '{{TRIP_NUMBER}}': data['trip_number'],
            '{{FLIGHT_TYPE}}': data['flight_type'],
            '{{DEPARTURE_DATE}}': data['departure_date'],
            '{{DEPARTURE_TIME_UTC}}': data['departure_time_utc'],
            '{{ARRIVAL_DATE}}': data['arrival_date'],
            '{{ARRIVAL_TIME_UTC}}': data['arrival_time_utc'],
            '{{FLIGHT_TIME}}': data['flight_time'],
            '{{DISTANCE}}': data['distance'],
            
            '{{ORIGIN_AIRPORT}}': data['origin_airport_name'],
            '{{ORIGIN_CODE}}': data['origin_airport_code'],
            '{{ORIGIN_ICAO}}': data['origin_airport_icao'],
            '{{ORIGIN_CITY}}': data['origin_city'],
            '{{ORIGIN_COUNTRY}}': data['origin_country'],
            
            '{{DESTINATION_AIRPORT}}': data['destination_airport_name'],
            '{{DESTINATION_CODE}}': data['destination_airport_code'],
            '{{DESTINATION_ICAO}}': data['destination_airport_icao'],
            '{{DESTINATION_CITY}}': data['destination_city'],
            '{{DESTINATION_COUNTRY}}': data['destination_country'],
            
            '{{AIRCRAFT_TAIL}}': data['aircraft']['tail_number'],
            '{{AIRCRAFT_MAKE}}': data['aircraft']['make'],
            '{{AIRCRAFT_MODEL}}': data['aircraft']['model'],
            '{{AIRCRAFT_COMPANY}}': data['aircraft']['company'],
            '{{AIRCRAFT_SERIAL}}': data['aircraft']['serial_number'],
            '{{AIRCRAFT_MGTOW}}': data['aircraft']['mgtow'],
            
            '{{PRIMARY_PILOT}}': data['crew']['primary_pilot'],
            '{{SECONDARY_PILOT}}': data['crew']['secondary_pilot'],
            '{{MEDICS}}': ', '.join(data['crew']['medics']) if data['crew']['medics'] else 'N/A',
            
            '{{PASSENGER_COUNT}}': str(data['passenger_count']),
            '{{TOTAL_PAX_COUNT}}': str(data['total_pax_count']),
            '{{IS_PAX_LEG}}': 'Yes' if data['is_pax_leg'] else 'No',
            '{{DOCUMENT_DATE}}': data['document_date'],
            '{{DOCUMENT_TIME}}': data['document_time'],
            
            # Official CBP Form 7507 required fields
            '{{CARRIER_CODE}}': data.get('carrier_code', 'N/A'),
            '{{FLIGHT_NUMBER}}': data.get('flight_number', data['trip_number']),
            '{{CREW_COUNT}}': str(data.get('crew_count', 2)),  # Pilot + Co-pilot minimum
            '{{PILOT_TITLE}}': 'Aircraft Commander',
        }
        
        # Add combined passenger and patient list (PAX leg logic)
        if data['passengers_and_patient']:
            pax_list = '\n'.join([
                f"{p['name']} ({p['type']}) - {p['nationality']} - Passport: {p['passport_number']}"
                for p in data['passengers_and_patient']
            ])
            replacements['{{PASSENGERS_AND_PATIENT_LIST}}'] = pax_list
        else:
            replacements['{{PASSENGERS_AND_PATIENT_LIST}}'] = 'No passengers or patient (Repositioning flight)'
        
        # Add legacy passenger list for backward compatibility
        if data['passengers']:
            passenger_list = '\n'.join([
                f"{p['name']} - {p['nationality']} - Passport: {p['passport_number']}"
                for p in data['passengers']
            ])
            replacements['{{PASSENGER_LIST}}'] = passenger_list
        else:
            replacements['{{PASSENGER_LIST}}'] = 'No passengers (Repositioning flight)'
        
        # Add patient info if available
        if data['patient']:
            replacements['{{PATIENT_NAME}}'] = data['patient']['name']
            replacements['{{PATIENT_DOB}}'] = data['patient']['date_of_birth']
        else:
            replacements['{{PATIENT_NAME}}'] = 'N/A'
            replacements['{{PATIENT_DOB}}'] = 'N/A'
        
        # Apply replacements
        full_text = paragraph.text
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # Update paragraph text if changes were made
        if full_text != paragraph.text:
            paragraph.text = full_text
    
    def _clean_name(self, name: str) -> str:

        if not name:
            return name
            
        # Common titles to remove (order matters - longer phrases first)
        titles = [
            'First Officer', 'F/O', 'FO',
            'Captain', 'Capt.', 'Capt',
            'Dr.', 'Dr', 'Doctor',
            'Nurse', 'RN', 'LPN',
            'Mr.', 'Mrs.', 'Ms.', 'Miss',
            'Sir', 'Madam', 'Ma\'am'
        ]
        
        # Handle compound titles like "First Officer"
        clean_name = name
        for title in titles:
            # Remove title from beginning with space
            if clean_name.startswith(title + ' '):
                clean_name = clean_name[len(title + ' '):]
                break
            # Remove title if it's the whole beginning part
            elif clean_name.startswith(title):
                clean_name = clean_name[len(title):].lstrip()
                break
        
        # Split into parts and clean up any remaining single-word titles
        parts = clean_name.split()
        
        # Join remaining parts (should be first and last name)
        return ' '.join(parts)
    
    def _fill_gendec_pdf(self, trip_line, template_path: Path) -> dict:

        import io
        
        # PDF form field mapping - Basic fields
        field_mapping = {
            'owner_operator': 'F[0].P1[0].OwnerorOperator[0]',
            'aircraft_registration': 'F[0].P1[0].MarksofNationality[0]',
            'flight_number': 'F[0].P1[0].flightnumber[0]',
            'date': 'F[0].P1[0].Date[0]',
            'departure_from': 'F[0].P1[0].departurefrom[0]',
            'arrive_at': 'F[0].P1[0].arrivaat[0]',
            'embarking': 'F[0].P1[0].embarking[0]',
            'through_passengers': 'F[0].P1[0].throughonsameflight[0]',
            'disembarking': 'F[0].P1[0].disembarking[0]',
            'through_passengers2': 'F[0].P1[0].throughonsameflight2[0]',
            'declaration': 'F[0].P1[0].declaration1[0]',
            'other_info': 'F[0].P1[0].other1[0]',
            'details': 'F[0].P1[0].details1[0]',
            'captain_signature': 'F[0].P1[0].sign[0]',
            'agent_signature': 'F[0].P1[0].agentsignature[0]',
            'awb_number': 'F[0].P1[0].AWB[0]'
        }
        
        # Crew and passenger place/total number field mapping
        place_fields = {
            'place1': 'F[0].P1[0].place1[0]',
            'place2': 'F[0].P1[0].place2[0]',
            'place3': 'F[0].P1[0].place3[0]',
            'place4': 'F[0].P1[0].place4[0]',
            'place5': 'F[0].P1[0].place5[0]',
            'place6': 'F[0].P1[0].place6[0]',
            'place7': 'F[0].P1[0].place7[0]',
            'place8': 'F[0].P1[0].place8[0]',
        }
        
        total_fields = {
            'total1': 'F[0].P1[0].totalnumber1[0]',
            'total2': 'F[0].P1[0].totalnumber2[0]',
            'total3': 'F[0].P1[0].totalnumber3[0]',
            'total4': 'F[0].P1[0].totalnumber4[0]',
            'total5': 'F[0].P1[0].totalnumber5[0]',
            'total6': 'F[0].P1[0].totalnumber6[0]',
            'total7': 'F[0].P1[0].totalnumber7[0]',
            'total8': 'F[0].P1[0].totalnumber8[0]',
        }
        
        # Prepare form data
        aircraft = trip_line.trip.aircraft
        
        # Format airports as CITY, COUNTRY (ICAO)
        def format_airport(airport):
            if not airport:
                return ''
            city = getattr(airport, 'city', airport.name)
            country = getattr(airport, 'country', 'US')  # Default to US if not available
            icao = getattr(airport, 'icao_code', '')
            return f"{city}, {country} ({icao})"
        
        data = {
            'owner_operator': getattr(aircraft, 'company', 'Unknown Operator'),
            'aircraft_registration': getattr(aircraft, 'tail_number', 'N-UNKNOWN'),
            'flight_number': getattr(trip_line.trip, 'flight_number', '') or f"TL{trip_line.id}",
            'date': trip_line.departure_time.strftime('%m/%d/%Y') if trip_line.departure_time else datetime.now().strftime('%m/%d/%Y'),
            'departure_from': format_airport(trip_line.origin_airport),
            'arrive_at': format_airport(trip_line.destination_airport),
        }
        
        # Prepare crew and passenger data for place/total fields
        crew_and_pax_data = {}
        field_index = 1
        
        # Add crew members first
        crew_line = getattr(trip_line, 'crew_line', None)
        if crew_line:
            # Primary in Command (PIC)
            if hasattr(crew_line, 'primary_in_command') and crew_line.primary_in_command:
                name = self._clean_name(crew_line.primary_in_command.name)
                crew_and_pax_data[f'place{field_index}'] = f"PIC: {name}"
                # Get birthdate from crew member info if available
                if hasattr(crew_line.primary_in_command, 'info') and hasattr(crew_line.primary_in_command.info, 'date_of_birth'):
                    crew_and_pax_data[f'total{field_index}'] = crew_line.primary_in_command.info.date_of_birth.strftime('%m/%d/%Y')
                else:
                    crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                field_index += 1
            
            # Second in Command (SIC)
            if hasattr(crew_line, 'secondary_in_command') and crew_line.secondary_in_command:
                name = self._clean_name(crew_line.secondary_in_command.name)
                crew_and_pax_data[f'place{field_index}'] = f"SIC: {name}"
                if hasattr(crew_line.secondary_in_command, 'info') and hasattr(crew_line.secondary_in_command.info, 'date_of_birth'):
                    crew_and_pax_data[f'total{field_index}'] = crew_line.secondary_in_command.info.date_of_birth.strftime('%m/%d/%Y')
                else:
                    crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                field_index += 1
            
            # Medical crew (MED)
            if hasattr(crew_line, 'medic_ids'):
                medics = list(crew_line.medic_ids.all()) if crew_line.medic_ids else []
                for medic in medics:
                    if field_index <= 8:  # Limit to available fields
                        name = self._clean_name(medic.name)
                        crew_and_pax_data[f'place{field_index}'] = f"MED: {name}"
                        if hasattr(medic, 'info') and hasattr(medic.info, 'date_of_birth'):
                            crew_and_pax_data[f'total{field_index}'] = medic.info.date_of_birth.strftime('%m/%d/%Y')
                        else:
                            crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                        field_index += 1
        
        # Add passengers if it's a PAX leg
        if trip_line.passenger_leg:
            # Add patient first
            patient = trip_line.trip.patient if hasattr(trip_line.trip, 'patient') and trip_line.trip.patient else None
            if patient and field_index <= 8:
                name = self._clean_name(patient.info.first_name + ' ' + patient.info.last_name if patient.info else patient.name)
                crew_and_pax_data[f'place{field_index}'] = f"PAX: {name}"
                if hasattr(patient, 'info') and patient.info and hasattr(patient.info, 'date_of_birth') and patient.info.date_of_birth:
                    crew_and_pax_data[f'total{field_index}'] = patient.info.date_of_birth.strftime('%m/%d/%Y')
                else:
                    crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                field_index += 1
            
            # Add regular passengers
            passengers = list(trip_line.trip.passengers.all()) if hasattr(trip_line.trip, 'passengers') else []
            for passenger in passengers:
                if field_index <= 8:  # Limit to available fields
                    # Use Contact's name and DOB (passenger.info is the Contact)
                    if passenger.info:
                        name = self._clean_name(f"{passenger.info.first_name or ''} {passenger.info.last_name or ''}".strip())
                        crew_and_pax_data[f'place{field_index}'] = f"PAX: {name}"
                        if passenger.info.date_of_birth:
                            crew_and_pax_data[f'total{field_index}'] = passenger.info.date_of_birth.strftime('%m/%d/%Y')
                        else:
                            crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                    else:
                        # Fallback to passenger name if no contact info
                        crew_and_pax_data[f'place{field_index}'] = f"PAX: Unknown"
                        crew_and_pax_data[f'total{field_index}'] = 'DOB not available'
                    field_index += 1
            
            # Calculate total passenger counts
            total_pax = len(passengers) + (1 if patient else 0)
            data.update({
                'embarking': str(total_pax),
                'disembarking': str(total_pax),
                'through_passengers': '0',
                'through_passengers2': '0',
            })
            
            # Generate passenger/patient details for details field
            details_parts = []
            if patient:
                details_parts.append(f"Patient: {patient.name}")
            if passengers:
                passenger_names = [p.name for p in passengers]
                details_parts.append(f"Passengers: {', '.join(passenger_names)}")
            data['details'] = ', '.join(details_parts) if details_parts else 'No passengers'
        else:
            data.update({
                'embarking': '0',
                'disembarking': '0',
                'through_passengers': '0',
                'through_passengers2': '0',
                'details': 'Repositioning flight - no passengers'
            })
        
        # Add crew and passenger data to main data dict
        data.update(crew_and_pax_data)
        
        # Add remaining fields
        primary_pilot = None
        if hasattr(trip_line.trip, 'crew_members'):
            pilots = [crew for crew in trip_line.trip.crew_members.all() if 'pilot' in crew.role.lower()]
            if pilots:
                primary_pilot = pilots[0].name
        
        data.update({
            'captain_signature': primary_pilot or 'Aircraft Commander',
            'agent_signature': data['owner_operator'],
            'declaration': 'This flight complies with all applicable US Customs regulations',
            'other_info': 'Medical transport flight' if trip_line.passenger_leg else 'Repositioning flight'
        })
        
        # Fill PDF form
        try:
            doc = fitz.open(template_path)
            # Combine all field mappings
            all_field_mappings = {**field_mapping, **place_fields, **total_fields}
        
            # Fill form fields
            filled_fields = 0
            total_fields = 0
        
            for page_num in range(doc.page_count):
                page = doc[page_num]
                widgets = list(page.widgets())
            
                for widget in widgets:
                    total_fields += 1
                    field_name = widget.field_name
                
                    # Find matching data key from all field mappings
                    data_key = None
                    for data_k, field_path in all_field_mappings.items():
                        if field_path == field_name:
                            data_key = data_k
                            break
                
                    if data_key and data_key in data:
                        widget.field_value = str(data[data_key])
                        widget.update()
                        filled_fields += 1
            
            # Save to memory buffer
            pdf_buffer = io.BytesIO()
            doc.save(pdf_buffer)
            pdf_buffer.seek(0)
            doc.close()
            
            return {
                'success': True,
                'pdf_data': pdf_buffer.getvalue(),
                'fields_filled': filled_fields,
                'total_fields': total_fields,
                'filename': f"GenDec_{trip_line.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            }
            
        except Exception as e:
            if 'doc' in locals():
                doc.close()
            raise Exception(f"PDF generation failed: {str(e)}")
    
    def _fill_template_fallback(self, template_path: Path, output_path: Path, data: Dict[str, Any]) -> str:

        # Simple copy of template as fallback
        import shutil
        shutil.copy2(template_path, output_path)
        
        print(f"Warning: Template filled using fallback method. Generated document: {output_path}")
        print("For full template processing, install python-docx: pip install python-docx")
        
        return str(output_path)

def generate_gendec_for_trip_line(trip_line_id: str) -> str:

    generator = DocumentGenerator()
    return generator.generate_general_declaration(trip_line_id)

def generate_gendec_for_trip(trip_id: str) -> List[str]:

    try:
        trip = Trip.objects.prefetch_related('trip_lines').get(id=trip_id)
    except Trip.DoesNotExist:
        raise Trip.DoesNotExist(f"Trip with ID {trip_id} not found")
    
    generator = DocumentGenerator()
    generated_docs = []
    
    for trip_line in trip.trip_lines.all():
        doc_path = generator.generate_general_declaration(str(trip_line.id))
        generated_docs.append(doc_path)
    
    return generated_docs

if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docgen.py <trip_line_id>")
        sys.exit(1)
    
    trip_line_id = sys.argv[1]
    
    try:
        doc_path = generate_gendec_for_trip_line(trip_line_id)
        print(f"GenDec generated successfully: {doc_path}")
    except Exception as e:
        print(f"Error generating GenDec: {e}")
        sys.exit(1)