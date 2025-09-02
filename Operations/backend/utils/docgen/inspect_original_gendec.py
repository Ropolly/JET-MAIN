#!/usr/bin/env python3

from pathlib import Path
from docx import Document

def inspect_original_template():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(" INSPECTING ORIGINAL GENDEC.DOCX TEMPLATE")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        print(" DOCUMENT PARAGRAPHS:")
        print("-" * 40)
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"[{i:2d}] {paragraph.text}")
        
        print(f"\n DOCUMENT TABLES ({len(doc.tables)} tables):")
        print("-" * 40)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n TABLE {table_idx + 1} ({len(table.rows)} rows x {len(table.columns)} cols):")
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(f"[{cell_idx}]: {cell_text}")
                    else:
                        row_data.append(f"[{cell_idx}]: (empty)")
                
                if row_data:
                    print(f"  Row {row_idx}: {' | '.join(row_data)}")
        
        # Look for form fields, content controls, or fillable areas
        print(f"\n LOOKING FOR FILLABLE FIELDS:")
        print("-" * 40)
        
        fillable_found = False
        
        # Check for underlines, brackets, or placeholder patterns
        all_text = []
        for paragraph in doc.paragraphs:
            all_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text.append(paragraph.text)
        
        document_text = ' '.join(all_text)
        
        # Look for common fillable field patterns
        patterns = ['_____', '{{', '}}', '[', ']', 'XXXX', '____']
        
        for pattern in patterns:
            if pattern in document_text:
                print(f" Found pattern '{pattern}' in document")
                fillable_found = True
        
        if not fillable_found:
            print("  No obvious fillable field patterns found")
            print(" Document may need manual field identification")
        
        print(f"\n DOCUMENT STATISTICS:")
        print(f"   • Total paragraphs: {len(doc.paragraphs)}")
        print(f"   • Total tables: {len(doc.tables)}")
        print(f"   • File size: {template_path.stat().st_size:,} bytes")
        
    except Exception as e:
        print(f" Error inspecting template: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    inspect_original_template()
