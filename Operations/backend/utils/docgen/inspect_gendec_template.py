#!/usr/bin/env python3

from pathlib import Path
from docx import Document

def inspect_template():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(f" Inspecting GenDec template: {template_path}")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print(f"Document has {len(doc.tables)} tables")
        print(f"Document has {len(doc.sections)} sections")
        print()
        
        # Extract all text from paragraphs
        print(" PARAGRAPHS:")
        print("-" * 20)
        placeholders_found = set()
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"P{i:2d}: {text}")
                
                # Look for placeholders
                import re
                found_placeholders = re.findall(r'\{\{[^}]+\}\}', text)
                placeholders_found.update(found_placeholders)
        
        print()
        
        # Extract text from tables
        print(" TABLE CONTENT:")
        print("-" * 20)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"Table {table_idx + 1}: {len(table.rows)} rows × {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(f"[{cell_text}]")
                        
                        # Look for placeholders in table cells
                        found_placeholders = re.findall(r'\{\{[^}]+\}\}', cell_text)
                        placeholders_found.update(found_placeholders)
                
                if row_text:
                    print(f"  Row {row_idx + 1}: {' | '.join(row_text)}")
            print()
        
        # Show all found placeholders
        print(" PLACEHOLDERS FOUND:")
        print("-" * 20)
        if placeholders_found:
            for placeholder in sorted(placeholders_found):
                print(f"  • {placeholder}")
        else:
            print("  No placeholders found!")
        
        print()
        
        # Show headers and footers
        print(" HEADERS & FOOTERS:")
        print("-" * 20)
        
        for section_idx, section in enumerate(doc.sections):
            print(f"Section {section_idx + 1}:")
            
            # Header
            header_text = []
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    header_text.append(paragraph.text.strip())
                    found_placeholders = re.findall(r'\{\{[^}]+\}\}', paragraph.text)
                    placeholders_found.update(found_placeholders)
            
            if header_text:
                print(f"  Header: {' | '.join(header_text)}")
            
            # Footer
            footer_text = []
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    footer_text.append(paragraph.text.strip())
                    found_placeholders = re.findall(r'\{\{[^}]+\}\}', paragraph.text)
                    placeholders_found.update(found_placeholders)
            
            if footer_text:
                print(f"  Footer: {' | '.join(footer_text)}")
        
        print()
        print(f" Template inspection complete")
        print(f" Total unique placeholders found: {len(placeholders_found)}")
        
    except Exception as e:
        print(f" Error reading template: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    inspect_template()
