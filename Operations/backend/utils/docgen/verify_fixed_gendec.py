#!/usr/bin/env python3

from pathlib import Path
from docx import Document

def verify_fixed_template():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(" VERIFYING FIXED GENDEC.DOCX TEMPLATE")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        # Extract all text from document
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        print(f" Document loaded successfully")
        print(f" Paragraphs: {len(doc.paragraphs)}")
        print(f" Tables: {len(doc.tables)}")
        print(f" File size: {template_path.stat().st_size:,} bytes")
        
        # Count placeholders
        placeholder_count = all_text.count('{{')
        print(f" Placeholder markers found: {placeholder_count}")
        
        # Look for critical CBP Form 7507 placeholders
        critical_placeholders = [
            '{{AIRCRAFT_COMPANY}}',
            '{{CARRIER_CODE}}', 
            '{{AIRCRAFT_TAIL}}',
            '{{AIRCRAFT_MAKE}}',
            '{{FLIGHT_NUMBER}}',
            '{{DEPARTURE_DATE}}',
            '{{ORIGIN_AIRPORT_NAME}}',
            '{{DEPARTURE_TIME_UTC}}',
            '{{DESTINATION_AIRPORT_NAME}}',
            '{{ARRIVAL_TIME_UTC}}',
            '{{CREW_COUNT}}',
            '{{TOTAL_PAX_COUNT}}',
            '{{PRIMARY_PILOT}}',
            '{{SECONDARY_PILOT}}',
            '{{PASSENGERS_AND_PATIENT_LIST}}',
            '{{DOCUMENT_DATE}}'
        ]
        
        print(f"\n CRITICAL PLACEHOLDER VERIFICATION:")
        print("-" * 50)
        
        found_placeholders = []
        missing_placeholders = []
        
        for placeholder in critical_placeholders:
            if placeholder in all_text:
                print(f" {placeholder}")
                found_placeholders.append(placeholder)
            else:
                print(f" {placeholder}")
                missing_placeholders.append(placeholder)
        
        print(f"\n SUMMARY:")
        print(f"   • Critical placeholders found: {len(found_placeholders)}/{len(critical_placeholders)}")
        print(f"   • Completion rate: {len(found_placeholders)/len(critical_placeholders)*100:.1f}%")
        
        if missing_placeholders:
            print(f"\n  MISSING PLACEHOLDERS:")
            for placeholder in missing_placeholders:
                print(f"   • {placeholder}")
            print(f"\n The template may need additional placeholders for full compliance")
        else:
            print(f"\n ALL CRITICAL PLACEHOLDERS PRESENT!")
            print(f" Template is ready for CBP Form 7507 generation")
        
        # Show first few lines of content to verify structure
        print(f"\n DOCUMENT PREVIEW (first 10 lines):")
        print("-" * 50)
        lines = all_text.split('\n')[:10]
        for i, line in enumerate(lines, 1):
            if line.strip():
                print(f"   {i:2d}: {line[:80]}{'...' if len(line) > 80 else ''}")
        
        # Check if it has both page 1 (form) and page 2 (instructions)
        has_form_content = 'Owner/Operator' in all_text or 'Aircraft Registration' in all_text
        has_instructions = 'Notes and Specifications' in all_text
        
        print(f"\n DOCUMENT STRUCTURE:")
        print(f"   • Has form content (Page 1): {'' if has_form_content else ''}")
        print(f"   • Has instructions (Page 2): {'' if has_instructions else ''}")
        
        if has_form_content and has_instructions:
            print(f"    Complete 2-page CBP Form 7507 structure")
        else:
            print(f"     Template may be incomplete")
        
    except Exception as e:
        print(f" Error verifying template: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    verify_fixed_template()
