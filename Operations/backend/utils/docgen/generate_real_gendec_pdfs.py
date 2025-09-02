#!/usr/bin/env python3

import subprocess
import tempfile
from pathlib import Path
import shutil
import sys
from datetime import datetime
from docx import Document
import os

def create_pdf_generator():

    class GenDecPDFGenerator:
        def __init__(self):
            self.base_dir = Path(__file__).parent
            self.templates_dir = self.base_dir / "documents"
            self.outputs_dir = self.base_dir / "pdf_outputs"
            self.outputs_dir.mkdir(exist_ok=True)
            
            # Check if LibreOffice is available
            self.libreoffice_available = self._check_libreoffice()
        
        def _check_libreoffice(self):

            try:
                result = subprocess.run(['libreoffice', '--version'], 
                                      capture_output=True, text=True, timeout=10)
                return result.returncode == 0
            except:
                return False
        
        def generate_gendec_pdf(self, data):

            # Use the fixed original GenDec.docx template (now contains form fields with placeholders)
            template_path = self.templates_dir / "GenDec.docx"
            if not template_path.exists():
                raise FileNotFoundError(f"GenDec template not found: {template_path}")
            print(f" Using fixed original GenDec template: {template_path.name}")
            
            # Create output filenames
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            flight_type = "PAX" if data.get('is_pax_leg', False) else "REPO"
            flight_num = data.get('flight_number', 'UNKNOWN')
            
            docx_filename = f"GenDec_{flight_type}_{flight_num}_{timestamp}.docx"
            pdf_filename = f"GenDec_{flight_type}_{flight_num}_{timestamp}.pdf"
            
            docx_path = self.outputs_dir / docx_filename
            pdf_path = self.outputs_dir / pdf_filename
            
            # Step 1: Fill DOCX template with data
            self._fill_docx_template(template_path, docx_path, data)
            
            # Step 2: Convert DOCX to PDF
            if self.libreoffice_available:
                success = self._convert_docx_to_pdf(docx_path, pdf_path)
                if success:
                    print(f" PDF generated: {pdf_filename}")
                    return str(pdf_path)
                else:
                    print(f"  PDF conversion failed, DOCX available: {docx_filename}")
                    return str(docx_path)
            else:
                print(f"  LibreOffice not available, DOCX generated: {docx_filename}")
                return str(docx_path)
        
        def _fill_docx_template(self, template_path, output_path, data):

            # Load template
            doc = Document(template_path)
            
            # Prepare replacements
            replacements = self._create_replacements(data)
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Replace in headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, replacements)
                
                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Save filled document
            doc.save(output_path)
            print(f" GenDec DOCX generated: {output_path.name}")
        
        def _convert_docx_to_pdf(self, docx_path, pdf_path):

            try:
                # Use LibreOffice to convert DOCX to PDF
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(self.outputs_dir),
                    str(docx_path)
                ]
                
                print(f" Converting {docx_path.name} to PDF...")
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF with same name as DOCX but .pdf extension
                    generated_pdf = self.outputs_dir / (docx_path.stem + '.pdf')
                    
                    if generated_pdf.exists():
                        # Rename to our desired filename if different
                        if generated_pdf != pdf_path:
                            generated_pdf.rename(pdf_path)
                        
                        file_size = pdf_path.stat().st_size
                        print(f" PDF conversion successful: {pdf_path.name} ({file_size:,} bytes)")
                        return True
                    else:
                        print(f" PDF file not found after conversion")
                        return False
                else:
                    print(f" LibreOffice conversion failed:")
                    print(f"   stdout: {result.stdout}")
                    print(f"   stderr: {result.stderr}")
                    return False
            
            except subprocess.TimeoutExpired:
                print(f" PDF conversion timeout")
                return False
            except Exception as e:
                print(f" PDF conversion error: {str(e)}")
                return False
        
        def _create_replacements(self, data):

            # Format passengers and patient list
            passenger_list = ""
            if data.get('passengers_and_patient'):
                lines = []
                for i, person in enumerate(data['passengers_and_patient'], 1):
                    line = f"{i}. {person['name']} ({person['type']}) - {person['nationality']} - Passport: {person['passport_number']} - DOB: {person['date_of_birth']}"
                    lines.append(line)
                passenger_list = '\n'.join(lines)
            else:
                passenger_list = "No passengers or patient (Repositioning flight)"
            
            # Create replacements dictionary
            replacements = {
                '{{TRIP_NUMBER}}': data.get('trip_number', 'N/A'),
                '{{DEPARTURE_DATE}}': data.get('departure_date', 'N/A'),
                '{{DEPARTURE_TIME_UTC}}': data.get('departure_time_utc', 'N/A'),
                '{{ARRIVAL_DATE}}': data.get('arrival_date', 'N/A'),
                '{{ARRIVAL_TIME_UTC}}': data.get('arrival_time_utc', 'N/A'),
                '{{ORIGIN_AIRPORT_NAME}}': data.get('origin_airport_name', 'N/A'),
                '{{ORIGIN_AIRPORT_CODE}}': data.get('origin_airport_code', 'N/A'),
                '{{ORIGIN_AIRPORT_ICAO}}': data.get('origin_airport_icao', 'N/A'),
                '{{ORIGIN_CITY}}': data.get('origin_city', 'N/A'),
                '{{DESTINATION_AIRPORT_NAME}}': data.get('destination_airport_name', 'N/A'),
                '{{DESTINATION_AIRPORT_CODE}}': data.get('destination_airport_code', 'N/A'),
                '{{DESTINATION_AIRPORT_ICAO}}': data.get('destination_airport_icao', 'N/A'),
                '{{DESTINATION_CITY}}': data.get('destination_city', 'N/A'),
                '{{AIRCRAFT_TAIL}}': data.get('aircraft', {}).get('tail_number', 'N/A'),
                '{{AIRCRAFT_MAKE}}': data.get('aircraft', {}).get('make', 'N/A'),
                '{{AIRCRAFT_MODEL}}': data.get('aircraft', {}).get('model', 'N/A'),
                '{{AIRCRAFT_COMPANY}}': data.get('aircraft', {}).get('company', 'N/A'),
                '{{AIRCRAFT_SERIAL}}': data.get('aircraft', {}).get('serial_number', 'N/A'),
                '{{AIRCRAFT_MGTOW}}': data.get('aircraft', {}).get('mgtow', 'N/A'),
                '{{PRIMARY_PILOT}}': data.get('crew', {}).get('primary_pilot', 'N/A'),
                '{{SECONDARY_PILOT}}': data.get('crew', {}).get('secondary_pilot', 'N/A'),
                '{{CREW_COUNT}}': str(data.get('crew_count', 0)),
                '{{PASSENGER_COUNT}}': str(data.get('passenger_count', 0)),
                '{{TOTAL_PAX_COUNT}}': str(data.get('total_pax_count', 0)),
                '{{CARRIER_CODE}}': data.get('carrier_code', 'N/A'),
                '{{FLIGHT_NUMBER}}': data.get('flight_number', data.get('trip_number', 'N/A')),
                '{{PILOT_TITLE}}': data.get('pilot_title', 'Aircraft Commander'),
                '{{PASSENGERS_AND_PATIENT_LIST}}': passenger_list,
                '{{DOCUMENT_DATE}}': data.get('document_date', datetime.now().strftime('%Y-%m-%d')),
                '{{DOCUMENT_TIME}}': data.get('document_time', datetime.now().strftime('%H:%M:%S UTC')),
            }
            
            return replacements
        
        def _replace_placeholders_in_paragraph(self, paragraph, replacements):

            full_text = paragraph.text
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
            
            if full_text != paragraph.text:
                paragraph.text = full_text
    
    return GenDecPDFGenerator()

def create_mock_pax_data():

    return {
        'trip_number': 'JET-2024-0902',
        'departure_date': '09/02/2024',
        'departure_time_utc': '14:30:00',
        'arrival_date': '09/02/2024',
        'arrival_time_utc': '18:45:00',
        
        'origin_airport_name': 'Teterboro Airport',
        'origin_airport_code': 'TEB',
        'origin_airport_icao': 'KTEB',
        'origin_city': 'Teterboro, NJ',
        
        'destination_airport_name': 'Miami International Airport',
        'destination_airport_code': 'MIA',
        'destination_airport_icao': 'KMIA',
        'destination_city': 'Miami, FL',
        
        'aircraft': {
            'tail_number': 'N525JT',
            'make': 'Cessna',
            'model': 'Citation CJ3+',
            'company': 'JET Charter Services LLC',
            'serial_number': '525C-0924',
            'mgtow': '13,870'
        },
        
        'crew': {
            'primary_pilot': 'Captain John Smith',
            'secondary_pilot': 'First Officer Sarah Johnson',
            'medics': ['Dr. Michael Brown', 'Paramedic Lisa Davis']
        },
        
        'passengers_and_patient': [
            {
                'name': 'Robert Wilson',
                'type': 'Patient',
                'nationality': 'United States',
                'passport_number': 'P123456789',
                'date_of_birth': '1965-03-15'
            },
            {
                'name': 'Emily Wilson',
                'type': 'Passenger',
                'nationality': 'United States',
                'passport_number': 'P987654321',
                'date_of_birth': '1968-07-22'
            },
            {
                'name': 'Dr. Jennifer Thompson',
                'type': 'Passenger',
                'nationality': 'United States',
                'passport_number': 'P456789123',
                'date_of_birth': '1972-11-08'
            }
        ],
        
        'passenger_count': 2,
        'total_pax_count': 3,
        'is_pax_leg': True,
        'crew_count': 4,
        'carrier_code': 'JCS',
        'flight_number': 'JET001',
        'pilot_title': 'Aircraft Commander'
    }

def create_mock_repo_data():

    return {
        'trip_number': 'JET-2024-0902R',
        'departure_date': '09/02/2024',
        'departure_time_utc': '20:15:00',
        'arrival_date': '09/02/2024',
        'arrival_time_utc': '20:45:00',
        
        'origin_airport_name': 'Miami International Airport',
        'origin_airport_code': 'MIA',
        'origin_airport_icao': 'KMIA',
        'origin_city': 'Miami, FL',
        
        'destination_airport_name': 'Fort Lauderdale Executive Airport',
        'destination_airport_code': 'FXE',
        'destination_airport_icao': 'KFXE',
        'destination_city': 'Fort Lauderdale, FL',
        
        'aircraft': {
            'tail_number': 'N525JT',
            'make': 'Cessna',
            'model': 'Citation CJ3+',
            'company': 'JET Charter Services LLC',
            'serial_number': '525C-0924',
            'mgtow': '13,870'
        },
        
        'crew': {
            'primary_pilot': 'Captain John Smith',
            'secondary_pilot': 'First Officer Sarah Johnson',
            'medics': []
        },
        
        'passengers_and_patient': [],
        
        'passenger_count': 0,
        'total_pax_count': 0,
        'is_pax_leg': False,
        'crew_count': 2,
        'carrier_code': 'JCS',
        'flight_number': 'JET001R',
        'pilot_title': 'Aircraft Commander'
    }

def main():

    print(" GENERATING REAL GENDEC PDF FILES")
    print("Using GenDec.docx template + LibreOffice conversion")
    print("=" * 60)
    
    try:
        # Create generator
        generator = create_pdf_generator()
        
        if not generator.libreoffice_available:
            print("  LibreOffice not available - will generate DOCX files only")
        else:
            print(" LibreOffice available for PDF conversion")
        
        # Test PAX leg
        print("\n Generating PAX Leg GenDec...")
        pax_data = create_mock_pax_data()
        pax_result = generator.generate_gendec_pdf(pax_data)
        
        # Test repositioning leg
        print("\n Generating Repositioning Leg GenDec...")
        repo_data = create_mock_repo_data()
        repo_result = generator.generate_gendec_pdf(repo_data)
        
        # Summary
        output_dir = generator.outputs_dir
        print(f"\n Generated files in: {output_dir}")
        
        # List generated files
        docx_files = list(output_dir.glob("*.docx"))
        pdf_files = list(output_dir.glob("*.pdf"))
        
        if pdf_files:
            print(f"\n PDF Files:")
            for f in pdf_files:
                size_kb = f.stat().st_size / 1024
                print(f"   • {f.name} ({size_kb:.1f} KB)")
        
        if docx_files:
            print(f"\n DOCX Files:")
            for f in docx_files:
                size_kb = f.stat().st_size / 1024
                print(f"   • {f.name} ({size_kb:.1f} KB)")
        
        print(f"\n GenDec generation complete!")
        
        if pdf_files:
            print(f" Generated {len(pdf_files)} PDF files from GenDec template")
        else:
            print(f" Generated DOCX files - manually convert to PDF if needed")
        
    except Exception as e:
        print(f" Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
