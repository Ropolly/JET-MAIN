#!/usr/bin/env python
"""
Enhanced Document Generation System for JET Aviation Operations

This module extends the basic document generation to create all document types
for a trip and save them with proper naming conventions.
"""

import os
import sys
import django
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
import uuid
import shutil

# Setup Django environment for standalone script execution
sys.path.append(str(Path(__file__).parent.parent.parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

from django.conf import settings
from django.utils import timezone
from api.models import Trip, Document, TripLine
from .docgen import DocumentGenerator

# Import python-docx Document with alias to avoid conflict
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Define document output directory
DOCUMENTS_DIR = Path(settings.BASE_DIR) / 'documents'


class TripDocumentGenerator:
    """Enhanced document generator that creates all trip documents."""
    
    def __init__(self, trip_id: str, user=None):
        """Initialize with a trip ID and optional user for tracking."""
        try:
            self.trip = Trip.objects.select_related(
                'aircraft', 'patient', 'patient__info', 'quote'
            ).prefetch_related(
                'trip_lines', 'passengers', 'documents'
            ).get(id=trip_id)
        except Trip.DoesNotExist:
            raise Trip.DoesNotExist(f"Trip with ID {trip_id} not found")
        
        self.user = user
        self.base_generator = DocumentGenerator()
        
        # Ensure documents directory exists
        DOCUMENTS_DIR.mkdir(exist_ok=True)
    
    def generate_document(self, document_type: str) -> Optional[Document]:
        """Generate a specific document type for the trip.
        
        Args:
            document_type: Type of document to generate (from Document.DOCUMENT_TYPES)
            
        Returns:
            Document instance if generated, None if not applicable
        """
        # Map document types to generation methods
        generators = {
            'gendec': self._generate_gendec,
            'quote': self._generate_quote_form,
            'customer_itinerary': self._generate_customer_itinerary,
            'internal_itinerary': self._generate_internal_itinerary,
            'payment_agreement': self._generate_payment_agreement,
            'consent_transport': self._generate_consent_transport,
            'psa': self._generate_psa,
            'handling_request': self._generate_handling_request,
        }
        
        generator_func = generators.get(document_type)
        if not generator_func:
            raise ValueError(f"Unknown document type: {document_type}")
        
        return generator_func()
    
    def generate_all_documents(self) -> List[Document]:
        """Generate all applicable documents for the trip.
        
        Returns:
            List of Document instances that were generated
        """
        generated_docs = []
        
        # Determine which documents to generate based on trip type
        doc_types_to_generate = self._get_applicable_document_types()
        
        for doc_type in doc_types_to_generate:
            try:
                doc = self.generate_document(doc_type)
                if doc:
                    generated_docs.append(doc)
            except Exception as e:
                print(f"Error generating {doc_type}: {e}")
                continue
        
        return generated_docs
    
    def _get_applicable_document_types(self) -> List[str]:
        """Determine which document types apply to this trip.
        
        Returns:
            List of document type codes to generate
        """
        doc_types = []
        
        # GenDec is needed for all trips with trip lines
        if self.trip.trip_lines.exists():
            doc_types.append('gendec')
            doc_types.append('handling_request')
        
        # Quote form if there's an associated quote
        if self.trip.quote:
            doc_types.append('quote')
        
        # Itineraries for all trips
        doc_types.append('customer_itinerary')
        doc_types.append('internal_itinerary')
        
        # Medical trips need additional documents
        if self.trip.type == 'medical':
            doc_types.append('payment_agreement')
            doc_types.append('consent_transport')
            if self.trip.patient:
                doc_types.append('psa')
        
        return doc_types
    
    def _generate_filename(self, doc_type: str, extension: str = 'docx') -> str:
        """Generate a unique filename for a document.
        
        Args:
            doc_type: Type of document
            extension: File extension
            
        Returns:
            Filename string
        """
        trip_number = self.trip.trip_number or "UNKNOWN"
        unique_id = str(uuid.uuid4())[:8]
        timestamp = datetime.now().strftime("%Y%m%d")
        
        return f"{trip_number}-{doc_type}-{timestamp}-{unique_id}.{extension}"
    
    def _check_and_remove_existing(self, doc_type: str):
        """Check for existing document of same type and remove if exists.
        
        Args:
            doc_type: Type of document to check
        """
        existing_docs = self.trip.documents.filter(document_type=doc_type)
        for doc in existing_docs:
            # Delete file from filesystem if it exists
            if doc.file_path and Path(doc.file_path).exists():
                try:
                    Path(doc.file_path).unlink()
                except Exception as e:
                    print(f"Error deleting file {doc.file_path}: {e}")
            # Delete database record
            doc.delete()
    
    def _save_document_record(self, source_path: str, doc_type: str, filename: str) -> Document:
        """Save document to filesystem and create database record.
        
        Args:
            source_path: Path to generated document
            doc_type: Type of document
            filename: Desired filename
            
        Returns:
            Document instance
        """
        # Check and remove existing document of same type
        self._check_and_remove_existing(doc_type)
        
        # Copy file to documents directory
        dest_path = DOCUMENTS_DIR / filename
        shutil.copy2(source_path, dest_path)
        
        # Create Document record
        doc = Document.objects.create(
            filename=filename,
            file_path=str(dest_path),
            document_type=doc_type,
            trip=self.trip,
            created_by=self.user
        )
        
        return doc
    
    def _generate_gendec(self) -> Optional[Document]:
        """Generate General Declaration documents for all trip lines.
        
        Returns:
            Document instance or None if no trip lines
        """
        if not self.trip.trip_lines.exists():
            return None
        
        # For now, generate for the first trip line
        # In future, could generate for all lines and combine
        trip_line = self.trip.trip_lines.first()
        
        try:
            # Use existing generator
            temp_path = self.base_generator.generate_general_declaration(str(trip_line.id))
            
            # Save with our naming convention
            filename = self._generate_filename('gendec')
            doc = self._save_document_record(temp_path, 'gendec', filename)
            
            # Clean up temp file
            if Path(temp_path).exists():
                Path(temp_path).unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating GenDec: {e}")
            return None
    
    def _generate_quote_form(self) -> Optional[Document]:
        """Generate Quote Form document with trip and quote data.
        
        Returns:
            Document instance or None if not applicable
        """
        if not self.trip.quote:
            return None
        
        template_path = self._get_template_path("QuoteForm.docx")
        if not template_path.exists():
            print("Quote form template not found")
            return None
        
        try:
            # Prepare data for quote document
            trip_data = self._prepare_trip_data()
            quote_data = self._prepare_quote_data()
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, quote_data)
            
            # Generate filled document
            filename = self._generate_filename('quote')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'quote', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating quote form: {e}")
            # Fallback to copying template
            filename = self._generate_filename('quote')
            doc = self._save_document_record(str(template_path), 'quote', filename)
            return doc
    
    def _generate_customer_itinerary(self) -> Optional[Document]:
        """Generate Customer Itinerary document with trip timeline data.
        
        Returns:
            Document instance
        """
        # Use appropriate template (templated version preferred)
        template_path = self._get_template_path("CustomerItin.docx")
        if not template_path.exists():
            # Fallback to generic itinerary
            template_path = self._get_template_path("Itinerary.docx")
        
        if not template_path.exists():
            print("Customer itinerary template not found")
            return None
        
        try:
            # Prepare data for customer itinerary
            trip_data = self._prepare_trip_data()
            quote_data = self._prepare_quote_data() if self.trip.quote else None
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, quote_data)
            
            # Generate filled document
            filename = self._generate_filename('customer_itinerary')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'customer_itinerary', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating customer itinerary: {e}")
            # Fallback to copying template
            filename = self._generate_filename('customer_itinerary')
            doc = self._save_document_record(str(template_path), 'customer_itinerary', filename)
            return doc
    
    def _generate_internal_itinerary(self) -> Optional[Document]:
        """Generate Internal Itinerary document with trip timeline data.
        
        Returns:
            Document instance
        """
        # Use appropriate template (templated version preferred)
        template_path = self._get_template_path("itin.docx")
        if not template_path.exists():
            # Fallback to generic itinerary
            template_path = self._get_template_path("Itinerary.docx")
        
        if not template_path.exists():
            print("Internal itinerary template not found")
            return None
        
        try:
            # Prepare data for internal itinerary (includes crew info)
            trip_data = self._prepare_trip_data()
            quote_data = self._prepare_quote_data() if self.trip.quote else None
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, quote_data)
            
            # Generate filled document
            filename = self._generate_filename('internal_itinerary')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'internal_itinerary', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating internal itinerary: {e}")
            # Fallback to copying template
            filename = self._generate_filename('internal_itinerary')
            doc = self._save_document_record(str(template_path), 'internal_itinerary', filename)
            return doc
    
    def _generate_payment_agreement(self) -> Optional[Document]:
        """Generate Payment Agreement document with patient and trip data.
        
        Returns:
            Document instance
        """
        template_path = self._get_template_path("PaymentAgreement.docx")
        if not template_path.exists():
            print("Payment agreement template not found")
            return None
        
        try:
            # Prepare data for payment agreement
            trip_data = self._prepare_trip_data()
            patient_data = self._prepare_patient_data() if self.trip.patient else None
            quote_data = self._prepare_quote_data() if self.trip.quote else None
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, quote_data, patient_data)
            
            # Generate filled document
            filename = self._generate_filename('payment_agreement')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'payment_agreement', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating payment agreement: {e}")
            # Fallback to copying template
            filename = self._generate_filename('payment_agreement')
            doc = self._save_document_record(str(template_path), 'payment_agreement', filename)
            return doc
    
    def _generate_consent_transport(self) -> Optional[Document]:
        """Generate Consent for Transport document with patient data.
        
        Returns:
            Document instance
        """
        template_path = self._get_template_path("Consent For Transport.docx")
        if not template_path.exists():
            print("Consent for transport template not found")
            return None
        
        try:
            # Prepare data for consent transport
            trip_data = self._prepare_trip_data()
            patient_data = self._prepare_patient_data() if self.trip.patient else None
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, None, patient_data)
            
            # Generate filled document
            filename = self._generate_filename('consent_transport')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'consent_transport', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating consent transport: {e}")
            # Fallback to copying template
            filename = self._generate_filename('consent_transport')
            doc = self._save_document_record(str(template_path), 'consent_transport', filename)
            return doc
    
    def _generate_psa(self) -> Optional[Document]:
        """Generate Patient Service Agreement document with patient data.
        
        Returns:
            Document instance
        """
        if not self.trip.patient:
            return None
        
        template_path = self._get_template_path("PSA.docx")
        if not template_path.exists():
            print("PSA template not found")
            return None
        
        try:
            # Prepare data for PSA
            trip_data = self._prepare_trip_data()
            patient_data = self._prepare_patient_data()
            quote_data = self._prepare_quote_data() if self.trip.quote else None
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data, quote_data, patient_data)
            
            # Generate filled document
            filename = self._generate_filename('psa')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'psa', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating PSA: {e}")
            # Fallback to copying template
            filename = self._generate_filename('psa')
            doc = self._save_document_record(str(template_path), 'psa', filename)
            return doc
    
    def _generate_handling_request(self) -> Optional[Document]:
        """Generate Handling Request document with aircraft and flight data.
        
        Returns:
            Document instance
        """
        template_path = self._get_template_path("HandlingRequest.docx")
        if not template_path.exists():
            print("Handling request template not found")
            return None
        
        try:
            # Prepare data for handling request
            trip_data = self._prepare_trip_data()
            
            # Create document data with placeholder mappings
            document_data = self._create_placeholder_mappings(trip_data)
            
            # Generate filled document
            filename = self._generate_filename('handling_request')
            temp_output_path = DOCUMENTS_DIR / f"temp_{filename}"
            
            filled_template_path = self._fill_template_with_data(
                str(template_path), 
                str(temp_output_path), 
                document_data
            )
            
            # Save with our naming convention
            doc = self._save_document_record(filled_template_path, 'handling_request', filename)
            
            # Clean up temp file
            if Path(filled_template_path).exists() and filled_template_path != str(temp_output_path):
                Path(filled_template_path).unlink()
            if temp_output_path.exists():
                temp_output_path.unlink()
            
            return doc
        except Exception as e:
            print(f"Error generating handling request: {e}")
            # Fallback to copying template
            filename = self._generate_filename('handling_request')
            doc = self._save_document_record(str(template_path), 'handling_request', filename)
            return doc
    
    def _prepare_trip_data(self) -> Dict[str, Any]:
        """Prepare comprehensive trip data for document generation.
        
        Returns:
            Dict containing all trip-related data for document placeholders
        """
        # Get first trip line for flight details (most documents use first leg)
        first_trip_line = self.trip.trip_lines.first() if self.trip.trip_lines.exists() else None
        
        # Basic trip information
        trip_data = {
            'trip_number': self.trip.trip_number or "TBD",
            'flight_type': self.trip.type.title() if self.trip.type else "Charter",
            'document_date': datetime.now().strftime("%Y-%m-%d"),
            'document_time': datetime.now().strftime("%H:%M"),
        }
        
        # Flight timing and route information
        if first_trip_line:
            trip_data.update({
                'departure_date': first_trip_line.departure_time_local.strftime("%Y-%m-%d") if first_trip_line.departure_time_local else "TBD",
                'departure_time_local': first_trip_line.departure_time_local.strftime("%H:%M") if first_trip_line.departure_time_local else "TBD",
                'departure_time_utc': first_trip_line.departure_time_utc.strftime("%H:%M UTC") if first_trip_line.departure_time_utc else "TBD",
                'arrival_date': first_trip_line.arrival_time_local.strftime("%Y-%m-%d") if first_trip_line.arrival_time_local else "TBD",
                'arrival_time_local': first_trip_line.arrival_time_local.strftime("%H:%M") if first_trip_line.arrival_time_local else "TBD",
                'arrival_time_utc': first_trip_line.arrival_time_utc.strftime("%H:%M UTC") if first_trip_line.arrival_time_utc else "TBD",
                'flight_time': f"{first_trip_line.flight_time} minutes" if first_trip_line.flight_time else "TBD",
                'distance': f"{first_trip_line.distance} nm" if first_trip_line.distance else "TBD",
                
                # Origin airport information
                'origin_airport_name': first_trip_line.origin_airport.name if first_trip_line.origin_airport else "TBD",
                'origin_airport_code': first_trip_line.origin_airport.iata_code or first_trip_line.origin_airport.ident if first_trip_line.origin_airport else "TBD",
                'origin_airport_icao': first_trip_line.origin_airport.icao_code if first_trip_line.origin_airport else "TBD",
                'origin_city': first_trip_line.origin_airport.municipality if first_trip_line.origin_airport else "TBD",
                'origin_country': first_trip_line.origin_airport.iso_country if first_trip_line.origin_airport else "TBD",
                
                # Destination airport information
                'destination_airport_name': first_trip_line.destination_airport.name if first_trip_line.destination_airport else "TBD",
                'destination_airport_code': first_trip_line.destination_airport.iata_code or first_trip_line.destination_airport.ident if first_trip_line.destination_airport else "TBD",
                'destination_airport_icao': first_trip_line.destination_airport.icao_code if first_trip_line.destination_airport else "TBD",
                'destination_city': first_trip_line.destination_airport.municipality if first_trip_line.destination_airport else "TBD",
                'destination_country': first_trip_line.destination_airport.iso_country if first_trip_line.destination_airport else "TBD",
            })
        else:
            # Default values when no trip lines exist
            trip_data.update({
                'departure_date': "TBD", 'departure_time_local': "TBD", 'departure_time_utc': "TBD",
                'arrival_date': "TBD", 'arrival_time_local': "TBD", 'arrival_time_utc': "TBD",
                'flight_time': "TBD", 'distance': "TBD",
                'origin_airport_name': "TBD", 'origin_airport_code': "TBD", 'origin_airport_icao': "TBD",
                'origin_city': "TBD", 'origin_country': "TBD",
                'destination_airport_name': "TBD", 'destination_airport_code': "TBD", 'destination_airport_icao': "TBD",
                'destination_city': "TBD", 'destination_country': "TBD",
            })
        
        # Aircraft information
        aircraft_data = {}
        if self.trip.aircraft:
            aircraft_data = {
                'tail_number': self.trip.aircraft.tail_number or "N/A",
                'type': f"{self.trip.aircraft.make} {self.trip.aircraft.model}".strip() or "N/A",
                'manufacturer': self.trip.aircraft.make or "N/A",
                'model': self.trip.aircraft.model or "N/A",
                'year': "N/A",  # Not available in Aircraft model
                'make': self.trip.aircraft.make or "N/A",
                'company': getattr(self.trip.aircraft, 'company', "N/A") or "N/A",
                'serial_number': getattr(self.trip.aircraft, 'serial_number', "N/A") or "N/A",
                'mgtow': str(getattr(self.trip.aircraft, 'mgtow', "N/A")) or "N/A",
            }
        else:
            aircraft_data = {
                'tail_number': "N/A", 'type': "N/A", 'manufacturer': "N/A", 
                'model': "N/A", 'year': "N/A", 'make': "N/A", 'company': "N/A",
                'serial_number': "N/A", 'mgtow': "N/A"
            }
        
        # Crew information
        crew_data = {'primary_pilot': "N/A", 'secondary_pilot': "N/A", 'medics': []}
        if first_trip_line and first_trip_line.crew_line:
            crew_line = first_trip_line.crew_line
            if crew_line.primary_in_command:
                crew_data['primary_pilot'] = self._format_contact_name(crew_line.primary_in_command)
            if crew_line.secondary_in_command:
                crew_data['secondary_pilot'] = self._format_contact_name(crew_line.secondary_in_command)
            crew_data['medics'] = [self._format_contact_name(medic) for medic in crew_line.medic_ids.all()]
        
        # Passenger information
        passenger_count = self.trip.passengers.count()
        passengers_info = []
        for passenger in self.trip.passengers.all():
            passengers_info.append({
                'name': self._format_contact_name(passenger.info),
                'nationality': getattr(passenger.info, 'nationality', 'N/A') or 'N/A',
                'date_of_birth': getattr(passenger.info, 'date_of_birth', None),
                'passport_number': getattr(passenger.info, 'passport_number', 'N/A') or 'N/A',
            })
        
        return {
            **trip_data,
            'aircraft': aircraft_data,
            'crew': crew_data,
            'passenger_count': passenger_count,
            'passengers': passengers_info,
        }
    
    def _prepare_quote_data(self) -> Dict[str, Any]:
        """Prepare quote-specific data for document generation.
        
        Returns:
            Dict containing quote-related data
        """
        if not self.trip.quote:
            return {}
        
        quote = self.trip.quote
        quote_data = {
            'quote_id': str(quote.id)[:8],  # Short ID for display
            'quoted_amount': f"${float(quote.quoted_amount):,.2f}" if quote.quoted_amount else "$0.00",
            'status': quote.status.title() if quote.status else "Pending",
            'estimated_flight_time': f"{quote.estimated_flight_time} hours" if quote.estimated_flight_time else "TBD",
            'aircraft_type': quote.aircraft_type or "TBD",
            'medical_team': "Yes" if quote.medical_team else "No",
            'includes_grounds': "Yes" if quote.includes_grounds else "No",
        }
        
        # Customer information from quote
        if quote.customer:
            quote_data['customer_name'] = self._format_contact_name(quote.customer)
            quote_data['customer_email'] = quote.customer.email or "N/A"
            quote_data['customer_phone'] = quote.customer.phone or "N/A"
        else:
            quote_data.update({
                'customer_name': "N/A", 'customer_email': "N/A", 'customer_phone': "N/A"
            })
        
        return quote_data
    
    def _prepare_patient_data(self) -> Dict[str, Any]:
        """Prepare patient-specific data for medical documents.
        
        Returns:
            Dict containing patient-related data
        """
        if not self.trip.patient:
            return {}
        
        patient = self.trip.patient
        patient_info = patient.info if patient.info else None
        
        patient_data = {
            'patient_name': self._format_contact_name(patient_info) if patient_info else "N/A",
            'patient_dob': patient_info.date_of_birth.strftime("%Y-%m-%d") if patient_info and patient_info.date_of_birth else "N/A",
            'patient_nationality': getattr(patient_info, 'nationality', 'N/A') or 'N/A' if patient_info else "N/A",
            'bed_at_origin': "Yes" if patient.bed_at_origin else "No",
            'bed_at_destination': "Yes" if patient.bed_at_destination else "No",
            'special_instructions': patient.special_instructions or "None",
            'patient_status': patient.status.title() if patient.status else "Active",
        }
        
        if patient_info:
            # Format address from individual components
            address_components = [
                patient_info.address_line1 or "",
                patient_info.address_line2 or "",
                patient_info.city or "",
                patient_info.state or "",
                patient_info.zip or "",
                patient_info.country or ""
            ]
            formatted_address = ", ".join([comp for comp in address_components if comp])
            
            patient_data.update({
                'patient_email': patient_info.email or "N/A",
                'patient_phone': patient_info.phone or "N/A",
                'patient_address': formatted_address or "N/A",
                'passport_number': getattr(patient_info, 'passport_number', 'N/A') or 'N/A',
            })
        
        return patient_data
    
    def _format_contact_name(self, contact) -> str:
        """Format a contact's name for display in documents.
        
        Args:
            contact: Contact instance or None
            
        Returns:
            Formatted name string
        """
        if not contact:
            return "N/A"
        
        first_name = getattr(contact, 'first_name', '') or ""
        last_name = getattr(contact, 'last_name', '') or ""
        
        if first_name and last_name:
            return f"{first_name} {last_name}"
        elif first_name:
            return first_name
        elif last_name:
            return last_name
        elif hasattr(contact, 'email') and contact.email:
            return contact.email
        else:
            return "N/A"
    
    def _get_template_path(self, template_name: str) -> Path:
        """Get template path, checking templated directory first.
        
        Args:
            template_name: Name of the template file (e.g., 'CustomerItin.docx')
            
        Returns:
            Path to the template file (templated version if available, otherwise original)
        """
        # Check templated directory first
        templated_path = self.base_generator.templates_dir / "templated" / template_name
        if templated_path.exists():
            return templated_path
        
        # Fall back to original template
        original_path = self.base_generator.templates_dir / template_name
        return original_path
    
    def _fill_template_with_data(self, template_path: str, output_path: str, placeholder_mappings: Dict[str, str]) -> str:
        """Fill a Word document template with trip data using placeholder replacement.
        
        Args:
            template_path: Path to template file
            output_path: Path for output file  
            placeholder_mappings: Dictionary mapping placeholders to values
            
        Returns:
            Path to generated document
        """
        try:
            if DocxDocument is not None:
                # Use python-docx for template filling
                doc = DocxDocument(template_path)
                
                # Replace placeholders in paragraphs
                for paragraph in doc.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, placeholder_mappings)
                
                # Replace placeholders in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_placeholders_in_paragraph(paragraph, placeholder_mappings)
                
                # Save the document
                doc.save(output_path)
                return output_path
            else:
                # Fallback method without python-docx
                return self._copy_and_replace_text(template_path, output_path, placeholder_mappings)
                
        except Exception as e:
            print(f"Error filling template: {e}. Falling back to copy method.")
            # Ultimate fallback - just copy the template
            import shutil
            shutil.copy2(template_path, output_path)
            return output_path
    
    def _replace_placeholders_in_paragraph(self, paragraph, placeholder_mappings: Dict[str, str]):
        """Replace placeholders in a paragraph with actual values.
        
        Args:
            paragraph: Document paragraph object
            placeholder_mappings: Dictionary mapping placeholders to values
        """
        full_text = paragraph.text
        for placeholder, value in placeholder_mappings.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # Clear existing runs and set new text
        if full_text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(full_text)
    
    def _copy_and_replace_text(self, template_path: str, output_path: str, placeholder_mappings: Dict[str, str]) -> str:
        """Fallback method to copy template and attempt basic text replacement.
        
        Args:
            template_path: Path to template file
            output_path: Path for output file
            placeholder_mappings: Dictionary mapping placeholders to values
            
        Returns:
            Path to output file
        """
        import shutil
        shutil.copy2(template_path, output_path)
        print(f"Template copied to {output_path}. Placeholder replacement not available without python-docx.")
        return output_path
    
    def _create_placeholder_mappings(self, trip_data: Dict[str, Any], quote_data: Dict[str, Any] = None, patient_data: Dict[str, Any] = None) -> Dict[str, str]:
        """Create comprehensive placeholder mappings for document generation.
        
        Args:
            trip_data: Trip data dictionary
            quote_data: Optional quote data dictionary
            patient_data: Optional patient data dictionary
            
        Returns:
            Dictionary mapping template placeholders to actual values
        """
        # Base trip placeholders
        placeholders = {
            '{{TRIP_NUMBER}}': trip_data.get('trip_number', 'TBD'),
            '{{FLIGHT_TYPE}}': trip_data.get('flight_type', 'Charter'),
            '{{DEPARTURE_DATE}}': trip_data.get('departure_date', 'TBD'),
            '{{DEPARTURE_TIME_LOCAL}}': trip_data.get('departure_time_local', 'TBD'),
            '{{DEPARTURE_TIME_UTC}}': trip_data.get('departure_time_utc', 'TBD'),
            '{{ARRIVAL_DATE}}': trip_data.get('arrival_date', 'TBD'),
            '{{ARRIVAL_TIME_LOCAL}}': trip_data.get('arrival_time_local', 'TBD'),
            '{{ARRIVAL_TIME_UTC}}': trip_data.get('arrival_time_utc', 'TBD'),
            '{{FLIGHT_TIME}}': trip_data.get('flight_time', 'TBD'),
            '{{DISTANCE}}': trip_data.get('distance', 'TBD'),
            
            # Airport placeholders
            '{{ORIGIN_AIRPORT}}': trip_data.get('origin_airport_name', 'TBD'),
            '{{ORIGIN_CODE}}': trip_data.get('origin_airport_code', 'TBD'),
            '{{ORIGIN_ICAO}}': trip_data.get('origin_airport_icao', 'TBD'),
            '{{ORIGIN_CITY}}': trip_data.get('origin_city', 'TBD'),
            '{{ORIGIN_COUNTRY}}': trip_data.get('origin_country', 'TBD'),
            
            '{{DESTINATION_AIRPORT}}': trip_data.get('destination_airport_name', 'TBD'),
            '{{DESTINATION_CODE}}': trip_data.get('destination_airport_code', 'TBD'),
            '{{DESTINATION_ICAO}}': trip_data.get('destination_airport_icao', 'TBD'),
            '{{DESTINATION_CITY}}': trip_data.get('destination_city', 'TBD'),
            '{{DESTINATION_COUNTRY}}': trip_data.get('destination_country', 'TBD'),
            
            # Aircraft placeholders
            '{{AIRCRAFT_TAIL}}': trip_data['aircraft']['tail_number'],
            '{{AIRCRAFT_TYPE}}': trip_data['aircraft']['type'],
            '{{AIRCRAFT_MANUFACTURER}}': trip_data['aircraft']['manufacturer'],
            '{{AIRCRAFT_MODEL}}': trip_data['aircraft']['model'],
            '{{AIRCRAFT_YEAR}}': trip_data['aircraft']['year'],
            '{{AIRCRAFT_MAKE}}': trip_data['aircraft']['make'],
            
            # Crew placeholders
            '{{PRIMARY_PILOT}}': trip_data['crew']['primary_pilot'],
            '{{SECONDARY_PILOT}}': trip_data['crew']['secondary_pilot'],
            '{{MEDICS}}': ', '.join(trip_data['crew']['medics']) if trip_data['crew']['medics'] else 'N/A',
            
            # Passenger placeholders
            '{{PASSENGER_COUNT}}': str(trip_data.get('passenger_count', 0)),
            
            # Document metadata
            '{{DOCUMENT_DATE}}': trip_data.get('document_date', datetime.now().strftime("%Y-%m-%d")),
            '{{DOCUMENT_TIME}}': trip_data.get('document_time', datetime.now().strftime("%H:%M")),
            
            # Company information
            '{{COMPANY_NAME}}': 'JET ICU Medical Transport',
            '{{COMPANY_ADDRESS}}': '1511 N Westshore Blvd #650, Tampa, FL 33607',
            '{{COMPANY_PHONE}}': '(352) 796-2540',
            '{{COMPANY_EMAIL}}': 'info@jeticu.com',
        }
        
        # Add quote-specific placeholders
        if quote_data:
            placeholders.update({
                '{{QUOTE_ID}}': quote_data.get('quote_id', 'N/A'),
                '{{QUOTED_AMOUNT}}': quote_data.get('quoted_amount', '$0.00'),
                '{{QUOTE_STATUS}}': quote_data.get('status', 'Pending'),
                '{{ESTIMATED_FLIGHT_TIME}}': quote_data.get('estimated_flight_time', 'TBD'),
                '{{AIRCRAFT_TYPE_QUOTED}}': quote_data.get('aircraft_type', 'TBD'),
                '{{MEDICAL_TEAM}}': quote_data.get('medical_team', 'No'),
                '{{INCLUDES_GROUNDS}}': quote_data.get('includes_grounds', 'No'),
                '{{CUSTOMER_NAME}}': quote_data.get('customer_name', 'N/A'),
                '{{CUSTOMER_EMAIL}}': quote_data.get('customer_email', 'N/A'),
                '{{CUSTOMER_PHONE}}': quote_data.get('customer_phone', 'N/A'),
            })
        
        # Add patient-specific placeholders
        if patient_data:
            placeholders.update({
                '{{PATIENT_NAME}}': patient_data.get('patient_name', 'N/A'),
                '{{PATIENT_DOB}}': patient_data.get('patient_dob', 'N/A'),
                '{{PATIENT_NATIONALITY}}': patient_data.get('patient_nationality', 'N/A'),
                '{{PATIENT_EMAIL}}': patient_data.get('patient_email', 'N/A'),
                '{{PATIENT_PHONE}}': patient_data.get('patient_phone', 'N/A'),
                '{{PATIENT_ADDRESS}}': patient_data.get('patient_address', 'N/A'),
                '{{PATIENT_PASSPORT}}': patient_data.get('passport_number', 'N/A'),
                '{{BED_AT_ORIGIN}}': patient_data.get('bed_at_origin', 'No'),
                '{{BED_AT_DESTINATION}}': patient_data.get('bed_at_destination', 'No'),
                '{{SPECIAL_INSTRUCTIONS}}': patient_data.get('special_instructions', 'None'),
                '{{PATIENT_STATUS}}': patient_data.get('patient_status', 'Active'),
            })
        
        # Add lowercase variants for templated documents
        lowercase_mappings = {
            '{{patient_name}}': placeholders.get('{{PATIENT_NAME}}', 'N/A'),
            '{{trip_date}}': placeholders.get('{{DEPARTURE_DATE}}', 'N/A'),
            '{{trip_number}}': placeholders.get('{{TRIP_NUMBER}}', 'N/A'),
            '{{aircraft_tail}}': placeholders.get('{{AIRCRAFT_TAIL}}', 'N/A'),
            '{{primary_pilot}}': placeholders.get('{{PRIMARY_PILOT}}', 'N/A'),
            '{{origin_airport}}': placeholders.get('{{ORIGIN_AIRPORT}}', 'N/A'),
            '{{destination_airport}}': placeholders.get('{{DESTINATION_AIRPORT}}', 'N/A'),
            '{{trip_itinerary}}': f"{placeholders.get('{{ORIGIN_AIRPORT}}', 'N/A')} → {placeholders.get('{{DESTINATION_AIRPORT}}', 'N/A')}",
            '{{customer_name}}': placeholders.get('{{CUSTOMER_NAME}}', 'N/A'),
            '{{quoted_amount}}': placeholders.get('{{QUOTED_AMOUNT}}', 'N/A'),
            '{{company_name}}': placeholders.get('{{COMPANY_NAME}}', 'N/A'),
        }
        
        # Merge lowercase mappings with uppercase ones
        placeholders.update(lowercase_mappings)
        
        return placeholders