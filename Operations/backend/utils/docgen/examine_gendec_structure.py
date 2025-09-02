#!/usr/bin/env python3

from pathlib import Path
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
import re

def examine_document_structure():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(" DEEP EXAMINATION OF GENDEC.DOCX STRUCTURE")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        # Look for content controls, form fields, or fillable areas
        print(" Examining document elements...")
        
        def extract_text_from_element(element):

            if hasattr(element, 'text'):
                return element.text
            elif hasattr(element, '_element'):
                return element._element.text if hasattr(element._element, 'text') else ''
            return ''
        
        # Check all document elements in order
        print("\n ALL DOCUMENT CONTENT (in order):")
        print("-" * 50)
        
        element_count = 0
        for element in doc.element.body:
            element_count += 1
            
            if element.tag.endswith('p'):  # Paragraph
                p = Paragraph(element, doc)
                text = p.text.strip()
                if text:
                    print(f"[P{element_count:2d}] {text}")
                    
                    # Look for underlines, form patterns
                    if '_' in text or re.search(r'\b[A-Z\s]+:\s*$', text):
                        print(f"      Potential form field: '{text}'")
            
            elif element.tag.endswith('tbl'):  # Table
                table = Table(element, doc)
                print(f"[T{element_count:2d}] TABLE ({len(table.rows)}x{len(table.columns)})")
                
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(f"C{cell_idx}: {cell_text}")
                            
                            # Look for form-like patterns in cells
                            if ('____' in cell_text or 
                                re.search(r':\s*$', cell_text) or
                                re.search(r'\.{3,}', cell_text)):
                                print(f"      Potential form field in cell: '{cell_text}'")
                    
                    if row_text:
                        print(f"     R{row_idx}: {' | '.join(row_text)}")
        
        # Look for specific patterns that indicate fillable areas
        print(f"\n SEARCHING FOR FILLABLE PATTERNS:")
        print("-" * 50)
        
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        
        # Common form patterns
        form_patterns = [
            (r'Aircraft.*Registration.*:', 'Aircraft Registration'),
            (r'Flight.*Number.*:', 'Flight Number'),
            (r'Pilot.*Command.*:', 'Pilot in Command'),
            (r'Departure.*:', 'Departure'),
            (r'Arrival.*:', 'Arrival'),
            (r'Crew.*:', 'Crew'),
            (r'Passenger.*:', 'Passenger'),
            (r'____+', 'Underline blanks'),
            (r'\.{3,}', 'Dot leaders'),
            (r'\[\s*\]', 'Checkboxes'),
            (r'Date.*:', 'Date fields'),
            (r'Time.*:', 'Time fields')
        ]
        
        found_patterns = []
        for pattern, description in form_patterns:
            matches = re.findall(pattern, all_text, re.IGNORECASE)
            if matches:
                print(f" {description}: {len(matches)} matches")
                found_patterns.append((description, matches))
            else:
                print(f" {description}: No matches")
        
        # Try to identify the main form structure
        print(f"\n MAIN FORM STRUCTURE ANALYSIS:")
        print("-" * 50)
        
        # Look for the main form table (usually the largest table)
        if doc.tables:
            main_table = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
            print(f" Main table: {len(main_table.rows)} rows x {len(main_table.columns)} columns")
            
            # Examine each cell for form-like content
            for row_idx, row in enumerate(main_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text and len(cell_text) < 100:  # Skip long text blocks
                        if (':' in cell_text or 
                            cell_text.endswith('...') or
                            '____' in cell_text):
                            print(f"  [R{row_idx}C{col_idx}] Form field candidate: '{cell_text}'")
        
        print(f"\n SUMMARY:")
        print(f"   • Document elements: {element_count}")
        print(f"   • Paragraphs: {len(doc.paragraphs)}")
        print(f"   • Tables: {len(doc.tables)}")
        print(f"   • Form patterns found: {len(found_patterns)}")
        
        if found_patterns:
            print(f"   • This appears to be a fillable form")
        else:
            print(f"   • This appears to be a static document")
    
    except Exception as e:
        print(f" Error examining document: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    examine_document_structure()
