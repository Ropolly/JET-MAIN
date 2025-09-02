#!/usr/bin/env python3

from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import xml.etree.ElementTree as ET

def extract_all_content():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(" EXTRACTING ALL GENDEC.DOCX CONTENT")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        # Check document sections and pages
        print(f" DOCUMENT STRUCTURE:")
        print(f"   • Sections: {len(doc.sections)}")
        
        # Extract raw XML to look for form fields
        print(f"\n EXAMINING RAW XML STRUCTURE...")
        
        # Look at document.xml content
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        # Get the main document part
        doc_part = doc.part
        
        print(f"   • Document part: {doc_part}")
        print(f"   • Related parts: {len(doc_part.related_parts)}")
        
        # Look for form fields in the XML
        root = doc_part._element
        
        # Find all text content including form fields
        print(f"\n ALL TEXT CONTENT (including form fields):")
        print("-" * 50)
        
        # Walk through all XML elements
        def walk_xml(element, depth=0):
            indent = "  " * depth
            
            if element.text and element.text.strip():
                print(f"{indent}TEXT: {element.text.strip()}")
            
            # Look for form field elements
            if 'fldChar' in element.tag:
                print(f"{indent}FORM FIELD: {element.tag}")
            elif 'instrText' in element.tag:
                print(f"{indent}FIELD INSTRUCTION: {element.text}")
            elif 'fldSimple' in element.tag:
                print(f"{indent}SIMPLE FIELD: {element.attrib}")
            
            for child in element:
                walk_xml(child, depth + 1)
        
        # walk_xml(root)
        
        # Try a different approach - look for all runs and their properties
        print(f"\n DETAILED CONTENT EXTRACTION:")
        print("-" * 50)
        
        content_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(f"PARAGRAPH: {paragraph.text}")
                
                # Check runs for special formatting
                for run in paragraph.runs:
                    if run.text.strip():
                        formatting = []
                        if run.bold:
                            formatting.append("BOLD")
                        if run.italic:
                            formatting.append("ITALIC")
                        if run.underline:
                            formatting.append("UNDERLINE")
                        
                        if formatting:
                            content_parts.append(f"  RUN ({'/'.join(formatting)}): {run.text}")
        
        # Check all tables thoroughly
        for table_idx, table in enumerate(doc.tables):
            content_parts.append(f"TABLE {table_idx + 1}:")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        content_parts.append(f"  [{row_idx},{cell_idx}]: {cell.text.strip()}")
                        
                        # Check for form controls in cells
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Look for underlines or special formatting
                                if run.underline or ('____' in run.text):
                                    content_parts.append(f"    POTENTIAL FORM FIELD: {run.text}")
        
        # Print all content parts
        for part in content_parts:
            print(part)
        
        # Look for content controls
        print(f"\n SEARCHING FOR CONTENT CONTROLS:")
        print("-" * 50)
        
        # Search for Word content controls in the XML
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        # Look for content control elements
        sdt_elements = root.xpath('.//w:sdt', namespaces=namespaces)
        if sdt_elements:
            print(f" Found {len(sdt_elements)} content controls")
            for i, sdt in enumerate(sdt_elements):
                print(f"  Content Control {i+1}: {sdt}")
        else:
            print(" No content controls found")
        
        # Look for form fields
        form_fields = root.xpath('.//w:fldChar', namespaces=namespaces)
        if form_fields:
            print(f" Found {len(form_fields)} form fields")
        else:
            print(" No form fields found")
        
        # Check if this is a multi-page document by looking for page breaks
        page_breaks = root.xpath('.//w:br[@w:type="page"]', namespaces=namespaces)
        section_breaks = root.xpath('.//w:sectPr', namespaces=namespaces)
        
        print(f"\n PAGE STRUCTURE:")
        print(f"   • Page breaks: {len(page_breaks)}")
        print(f"   • Section breaks: {len(section_breaks)}")
        
        if len(page_breaks) > 0 or len(section_breaks) > 1:
            print(" This appears to be a multi-page document")
            print(" The actual form fields may be on page 1")
        
    except Exception as e:
        print(f" Error extracting content: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    extract_all_content()
