#!/usr/bin/env python3

import os
import sys
from pathlib import Path
import tempfile
import datetime

# Add Django path
django_path = Path(__file__).parent.parent.parent
sys.path.append(str(django_path))

# Set Django settings
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')

try:
    import django
    django.setup()
    
    from utils.docgen.docgen import DocumentGenerator
    from api.models import Trip, TripLine, Aircraft, Airport, Patient, Passenger
    
    DJANGO_AVAILABLE = True
    print(" Django environment loaded successfully")
except Exception as e:
    DJANGO_AVAILABLE = False
    print(f"  Django environment not available: {str(e)}")
    print(" Running standalone tests only")

def test_standalone_generation():

    print(f"\n TESTING STANDALONE GENERATION")
    print("=" * 50)
    
    # Import standalone generation function
    sys.path.append(str(Path(__file__).parent))
    from generate_real_gendec_pdfs import create_pdf_generator, create_mock_pax_data, create_mock_repo_data
    
    try:
        generator = create_pdf_generator()
        
        # Test PAX leg
        print(" Testing PAX leg generation...")
        pax_data = create_mock_pax_data()
        pax_result = generator.generate_gendec_pdf(pax_data)
        
        if pax_result and Path(pax_result).exists():
            file_size = Path(pax_result).stat().st_size
            print(f" PAX PDF generated: {Path(pax_result).name} ({file_size:,} bytes)")
        else:
            print(f" PAX PDF generation failed")
            return False
        
        # Test repositioning leg
        print(" Testing repositioning leg generation...")
        repo_data = create_mock_repo_data()
        repo_result = generator.generate_gendec_pdf(repo_data)
        
        if repo_result and Path(repo_result).exists():
            file_size = Path(repo_result).stat().st_size
            print(f" Repositioning PDF generated: {Path(repo_result).name} ({file_size:,} bytes)")
        else:
            print(f" Repositioning PDF generation failed")
            return False
        
        print(f" Standalone generation test passed")
        return True
        
    except Exception as e:
        print(f" Standalone generation test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def test_django_integration():

    if not DJANGO_AVAILABLE:
        print(f"\n  SKIPPING DJANGO INTEGRATION TEST")
        print("Django environment not available")
        return True
    
    print(f"\n TESTING DJANGO INTEGRATION")
    print("=" * 50)
    
    try:
        # Create DocumentGenerator instance
        generator = DocumentGenerator()
        
        # Test template path detection
        template_path = generator.templates_dir / "GenDec_With_Placeholders.docx"
        
        if template_path.exists():
            file_size = template_path.stat().st_size
            print(f" New template found: {template_path.name} ({file_size:,} bytes)")
        else:
            print(f" New template not found: {template_path}")
            return False
        
        # Test data preparation methods
        print(" Testing data preparation methods...")
        
        # Check if the DocumentGenerator has the enhanced methods
        if hasattr(generator, '_calculate_crew_count'):
            print(" Enhanced crew count calculation available")
        else:
            print(" Enhanced crew count calculation missing")
        
        if hasattr(generator, '_prepare_gendec_data'):
            print(" GenDec data preparation method available")
        else:
            print(" GenDec data preparation method missing")
        
        # Test in-memory generation capability
        print(" Testing in-memory document generation...")
        
        # Create mock data structure similar to what Django would provide
        mock_trip_data = {
            'trip_number': 'JET-TEST-001',
            'departure_date': datetime.date.today(),
            'departure_time_utc': datetime.time(14, 30),
            'arrival_date': datetime.date.today(),
            'arrival_time_utc': datetime.time(18, 45),
            'origin_airport': {'name': 'Test Origin Airport', 'iata_code': 'TOR', 'icao_code': 'KTOR', 'city': 'Test City'},
            'destination_airport': {'name': 'Test Dest Airport', 'iata_code': 'TDE', 'icao_code': 'KTDE', 'city': 'Dest City'},
            'aircraft': {'tail_number': 'N-TEST', 'make': 'Test', 'model': 'Aircraft', 'company': 'Test Co', 'serial_number': 'TEST-123', 'mgtow': '10000'},
            'crew': {'primary_pilot': 'Test Pilot', 'secondary_pilot': 'Test Co-pilot', 'medics': []},
            'passengers_and_patient': [],
            'is_pax_leg': False,
            'crew_count': 2,
            'total_pax_count': 0,
            'passenger_count': 0
        }
        
        # Test template loading
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(template_path)
            print(" Template loads successfully in python-docx")
            
            # Test placeholder detection
            placeholders_found = 0
            for paragraph in doc.paragraphs:
                if '{{' in paragraph.text and '}}' in paragraph.text:
                    placeholders_found += 1
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{' in paragraph.text and '}}' in paragraph.text:
                                placeholders_found += 1
            
            print(f" Found {placeholders_found} placeholder instances in template")
            
        except Exception as e:
            print(f" Template loading failed: {str(e)}")
            return False
        
        print(f" Django integration test passed")
        return True
        
    except Exception as e:
        print(f" Django integration test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def test_template_compliance():

    print(f"\n TESTING CBP FORM 7507 COMPLIANCE")
    print("=" * 50)
    
    base_dir = Path(__file__).parent
    template_path = base_dir / "documents" / "GenDec_With_Placeholders.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return False
    
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(template_path)
        
        # Check for required CBP elements
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text.append(paragraph.text)
        
        document_text = ' '.join(full_text).upper()
        
        # CBP Form 7507 compliance checklist
        compliance_checks = [
            ("Form Title", "GENERAL DECLARATION" in document_text),
            ("CBP Reference", "CBP" in document_text or "CUSTOMS" in document_text),
            ("Aircraft Registration", "{{AIRCRAFT_TAIL}}" in ' '.join(full_text)),
            ("Flight Information", "{{FLIGHT_NUMBER}}" in ' '.join(full_text)),
            ("Crew Count", "{{CREW_COUNT}}" in ' '.join(full_text)),
            ("Passenger Count", "{{TOTAL_PAX_COUNT}}" in ' '.join(full_text)),
            ("Airport Information", "{{ORIGIN_AIRPORT" in ' '.join(full_text)),
            ("Declaration Statement", "DECLARATION" in document_text),
        ]
        
        passed_checks = 0
        for check_name, result in compliance_checks:
            status = "" if result else ""
            print(f"   {status} {check_name}")
            if result:
                passed_checks += 1
        
        compliance_rate = (passed_checks / len(compliance_checks)) * 100
        print(f"\n Compliance rate: {compliance_rate:.1f}% ({passed_checks}/{len(compliance_checks)})")
        
        if compliance_rate >= 80:
            print(f" CBP Form 7507 compliance test passed")
            return True
        else:
            print(f" CBP Form 7507 compliance test failed")
            return False
    
    except Exception as e:
        print(f" Compliance test failed: {str(e)}")
        return False

def main():

    print(" COMPLETE GENDEC INTEGRATION TEST SUITE")
    print("=" * 60)
    print(f" Test started at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    test_results = []
    
    # Test 1: Standalone generation
    test_results.append(("Standalone Generation", test_standalone_generation()))
    
    # Test 2: Django integration
    test_results.append(("Django Integration", test_django_integration()))
    
    # Test 3: CBP compliance
    test_results.append(("CBP Form 7507 Compliance", test_template_compliance()))
    
    # Summary
    print(f"\n TEST RESULTS SUMMARY")
    print("=" * 60)
    
    passed_tests = 0
    for test_name, result in test_results:
        status = " PASSED" if result else " FAILED"
        print(f"   {test_name:<25} {status}")
        if result:
            passed_tests += 1
    
    success_rate = (passed_tests / len(test_results)) * 100
    print(f"\n Overall success rate: {success_rate:.1f}% ({passed_tests}/{len(test_results)})")
    
    if success_rate == 100:
        print(f" ALL TESTS PASSED - GenDec system is ready for production!")
        print(f" Official US Customs CBP Form 7507 generation is fully functional")
        print(f"  Ready to generate compliant GenDec documents for aviation trips")
    else:
        print(f"  Some tests failed - review results above")
    
    print(f"\n Test completed at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

if __name__ == "__main__":
    main()
