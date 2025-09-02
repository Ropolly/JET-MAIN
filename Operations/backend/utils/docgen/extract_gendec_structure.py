#!/usr/bin/env python3

from pathlib import Path
from docx import Document

def extract_detailed_structure():

    template_path = Path(__file__).parent / "documents" / "GenDec.docx"
    
    if not template_path.exists():
        print(f" Template not found: {template_path}")
        return
    
    print(f" Detailed GenDec template structure")
    print("=" * 60)
    
    try:
        doc = Document(template_path)
        
        # Show ALL paragraph content with more detail
        print(" ALL PARAGRAPHS (detailed):")
        print("-" * 40)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            # Show even empty paragraphs
            if text or i < 20:  # Show first 20 paragraphs regardless
                if text.strip():
                    print(f"P{i:2d}: '{text}'")
                else:
                    print(f"P{i:2d}: [EMPTY]")
        
        print()
        
        # Look at the table structure in more detail
        print(" TABLE STRUCTURE (detailed):")
        print("-" * 40)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx + 1}:")
            print(f"  Dimensions: {len(table.rows)} rows × {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                print(f"  Row {row_idx + 1}:")
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    print(f"    Cell [{row_idx+1},{cell_idx+1}]: '{cell_text}'")
        
        print()
        
        # Check if this might be the form fields section that we missed
        print(" LOOKING FOR FORM FIELDS:")
        print("-" * 30)
        
        # Check for any runs that might contain form fields
        for i, paragraph in enumerate(doc.paragraphs):
            if i < 15:  # Focus on first part where form should be
                print(f"Paragraph {i} runs:")
                for run_idx, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        print(f"  Run {run_idx}: '{run.text}'")
        
        print()
        print(" ANALYSIS:")
        print("-" * 15)
        print("The template appears to be the official CBP form without placeholders.")
        print("We need to create a modified version with placeholders for data filling.")
        
    except Exception as e:
        print(f" Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    extract_detailed_structure()
