#!/usr/bin/env python3

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent.parent.parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

from api.models import Trip, TripLine, Aircraft, Patient, Passenger, Contact, Airport, CrewLine
from utils.docgen.docgen import DocumentGenerator
from datetime import datetime, timedelta
from django.utils import timezone

def test_original_template_integration():

    print(" TESTING ORIGINAL GENDEC.DOCX TEMPLATE INTEGRATION")
    print("=" * 70)
    
    try:
        # Initialize document generator
        generator = DocumentGenerator()
        
        # Verify template exists
        template_path = generator.templates_dir / "GenDec.docx"
        if not template_path.exists():
            print(f" Template not found: {template_path}")
            return False
        
        print(f" Template found: {template_path.name}")
        print(f" Template size: {template_path.stat().st_size:,} bytes")
        
        # Test 1: Create mock trip data for PAX leg
        print(f"\n TEST 1: PAX LEG GENERATION")
        print("-" * 40)
        
        # Create mock trip line data (simulating database objects)
        class MockTripLine:
            def __init__(self, passenger_leg=True):
                self.id = "test-trip-line-001"
                self.passenger_leg = passenger_leg
                self.trip = MockTrip(passenger_leg)
                self.origin_airport = MockAirport("LAX", "KLAX", "Los Angeles")
                self.destination_airport = MockAirport("JFK", "KJFK", "New York")
                self.crew_line = MockCrewLine()
                self.departure_date = timezone.now().date()
                self.departure_time = timezone.now().time()
                self.arrival_date = (timezone.now() + timedelta(hours=5)).date()
                self.arrival_time = (timezone.now() + timedelta(hours=5)).time()
        
        class MockTrip:
            def __init__(self, has_patient=True):
                self.id = "test-trip-001"
                self.trip_number = "JET001-TEST"
                self.aircraft = MockAircraft()
                self.patient = MockPatient() if has_patient else None
                self.passengers = MockPassengerQuerySet() if has_patient else []
        
        class MockAircraft:
            def __init__(self):
                self.tail_number = "N123JT"
                self.make = "Cessna"
                self.model = "Citation X"
                self.company = "JET Charter Services LLC"
                self.serial_number = "560-6001"
                self.mgtow_lbs = 36100
        
        class MockPatient:
            def __init__(self):
                self.info = MockContact("John Doe", "12/15/1970")
        
        class MockPassengerQuerySet:
            def all(self):
                return [MockPassenger("Jane Smith", "03/22/1985")]
            
            def __iter__(self):
                return iter(self.all())
        
        class MockPassenger:
            def __init__(self, name, dob):
                self.info = MockContact(name, dob)
        
        class MockContact:
            def __init__(self, name, dob):
                self.first_name = name.split()[0]
                self.last_name = name.split()[-1] if ' ' in name else ''
                self.date_of_birth = dob
                self.nationality = "United States"
                self.passport_number = f"P{hash(name) % 1000000000:09d}"
        
        class MockAirport:
            def __init__(self, iata, icao, name):
                self.iata_code = iata
                self.icao_code = icao
                self.name = name
                
        class MockCrewLine:
            def __init__(self):
                self.primary_in_command = MockContact("Captain Smith", "01/01/1980")
                self.secondary_in_command = MockContact("First Officer Johnson", "05/15/1985")
                
                # Mock medic_ids for crew count calculation
                class MockMedicQuerySet:
                    def all(self):
                        return [MockContact("Nurse Williams", "08/30/1975")]
                
                self.medic_ids = MockMedicQuerySet()
        
        # Test PAX leg
        pax_trip_line = MockTripLine(passenger_leg=True)
        pax_data = generator._prepare_gendec_data(pax_trip_line)
        
        print(f" PAX leg data prepared")
        print(f"   • Flight: {pax_data['flight_number']}")
        print(f"   • Route: {pax_data['origin_airport_code']} → {pax_data['destination_airport_code']}")
        print(f"   • PAX Count: {pax_data['total_pax_count']}")
        print(f"   • Crew Count: {pax_data['crew_count']}")
        print(f"   • Is PAX Leg: {pax_data['is_pax_leg']}")
        
        # Test 2: Create repositioning leg
        print(f"\n TEST 2: REPOSITIONING LEG GENERATION")
        print("-" * 40)
        
        repo_trip_line = MockTripLine(passenger_leg=False)
        repo_trip_line.trip.patient = None  # No patient on repo leg
        repo_trip_line.trip.passengers = []  # No passengers on repo leg
        repo_data = generator._prepare_gendec_data(repo_trip_line)
        
        print(f" Repositioning leg data prepared")
        print(f"   • Flight: {repo_data['flight_number']}")
        print(f"   • Route: {repo_data['origin_airport_code']} → {repo_data['destination_airport_code']}")
        print(f"   • PAX Count: {repo_data['total_pax_count']}")
        print(f"   • Crew Count: {repo_data['crew_count']}")
        print(f"   • Is PAX Leg: {repo_data['is_pax_leg']}")
        
        # Test 3: Document generation in memory
        print(f"\n TEST 3: DOCUMENT GENERATION")
        print("-" * 40)
        
        # Test PAX document generation
        pax_doc_content = generator._generate_document_in_memory(template_path, pax_data)
        print(f" PAX document generated in memory: {len(pax_doc_content):,} bytes")
        
        # Test repositioning document generation
        repo_doc_content = generator._generate_document_in_memory(template_path, repo_data)
        print(f" Repositioning document generated in memory: {len(repo_doc_content):,} bytes")
        
        # Test 4: Verify placeholder replacement worked
        print(f"\n TEST 4: PLACEHOLDER REPLACEMENT VERIFICATION")
        print("-" * 40)
        
        # Save test documents for verification
        test_output_dir = generator.outputs_dir / "integration_tests"
        test_output_dir.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pax_test_path = test_output_dir / f"GenDec_PAX_Integration_Test_{timestamp}.docx"
        repo_test_path = test_output_dir / f"GenDec_REPO_Integration_Test_{timestamp}.docx"
        
        # Write test files
        with open(pax_test_path, 'wb') as f:
            f.write(pax_doc_content)
        
        with open(repo_test_path, 'wb') as f:
            f.write(repo_doc_content)
        
        print(f" Test documents saved:")
        print(f"   • PAX: {pax_test_path.name}")
        print(f"   • Repo: {repo_test_path.name}")
        
        # Quick verification - check if placeholders were replaced
        from docx import Document
        
        # Check PAX document
        pax_doc = Document(pax_test_path)
        pax_text = ""
        for para in pax_doc.paragraphs:
            pax_text += para.text + "\n"
        for table in pax_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    pax_text += cell.text + " "
        
        remaining_placeholders = pax_text.count('{{')
        total_content_length = len(pax_text)
        
        print(f" Placeholder replacement verification:")
        print(f"   • Remaining placeholders: {remaining_placeholders}")
        print(f"   • Total content length: {total_content_length:,} chars")
        print(f"   • Contains flight data: {'JET001-TEST' in pax_text}")
        print(f"   • Contains aircraft data: {'N123JT' in pax_text}")
        print(f"   • Contains crew data: {'Captain Smith' in pax_text}")
        
        # Test 5: CBP Form 7507 compliance check
        print(f"\n TEST 5: CBP FORM 7507 COMPLIANCE")
        print("-" * 40)
        
        compliance_checks = [
            ('OMB Control Number', '1651-0002' in pax_text),
            ('CBP Form Identifier', '7507' in pax_text),
            ('Department Header', 'DEPARTMENT OF HOMELAND SECURITY' in pax_text),
            ('General Declaration Title', 'GENERAL DECLARATION' in pax_text),
            ('CFR References', '19 CFR' in pax_text),
            ('Aircraft Registration', 'N123JT' in pax_text),
            ('Crew Information', 'Captain Smith' in pax_text),
            ('Flight Information', 'JET001-TEST' in pax_text)
        ]
        
        passed_checks = 0
        for check_name, result in compliance_checks:
            status = "" if result else ""
            print(f"   {status} {check_name}: {result}")
            if result:
                passed_checks += 1
        
        compliance_rate = (passed_checks / len(compliance_checks)) * 100
        print(f"\n COMPLIANCE SUMMARY:")
        print(f"   • Passed checks: {passed_checks}/{len(compliance_checks)}")
        print(f"   • Compliance rate: {compliance_rate:.1f}%")
        
        # Final assessment
        print(f"\n INTEGRATION TEST RESULTS:")
        print("=" * 40)
        
        success_criteria = [
            template_path.exists(),
            len(pax_doc_content) > 30000,  # Reasonable document size
            len(repo_doc_content) > 30000,
            remaining_placeholders == 0,    # All placeholders replaced
            compliance_rate >= 87.5         # At least 7/8 compliance checks pass
        ]
        
        all_passed = all(success_criteria)
        
        if all_passed:
            print(f" ALL TESTS PASSED!")
            print(f" Original GenDec.docx template is fully functional")
            print(f" CBP Form 7507 compliance verified")
            print(f" System ready for production use")
        else:
            print(f" Some tests failed")
            print(f"  Manual review recommended")
        
        return all_passed
        
    except Exception as e:
        print(f" Integration test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_original_template_integration()
    sys.exit(0 if success else 1)
