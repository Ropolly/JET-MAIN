#!/usr/bin/env python3

from pathlib import Path
from docx import Document
from datetime import datetime
import io

def test_standalone_template():

    print(" STANDALONE GENDEC.DOCX TEMPLATE TEST")
    print("=" * 60)
    
    # Setup paths
    base_dir = Path(__file__).parent
    template_path = base_dir / "documents" / "GenDec.docx"
    output_dir = base_dir / "outputs" / "standalone_tests"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # Test 1: Template loading
        print(" TEST 1: TEMPLATE LOADING")
        print("-" * 40)
        
        if not template_path.exists():
            print(f" Template not found: {template_path}")
            return False
        
        doc = Document(template_path)
        print(f" Template loaded successfully")
        print(f" File size: {template_path.stat().st_size:,} bytes")
        print(f" Paragraphs: {len(doc.paragraphs)}")
        print(f" Tables: {len(doc.tables)}")
        
        # Test 2: Extract template content
        print(f"\n TEST 2: CONTENT ANALYSIS")
        print("-" * 40)
        
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        placeholder_count = all_text.count('{{')
        print(f" Total content length: {len(all_text):,} characters")
        print(f" Placeholder markers found: {placeholder_count}")
        
        # Test 3: Verify critical placeholders
        print(f"\n TEST 3: PLACEHOLDER VERIFICATION")
        print("-" * 40)
        
        critical_placeholders = [
            '{{AIRCRAFT_COMPANY}}',
            '{{AIRCRAFT_TAIL}}',
            '{{FLIGHT_NUMBER}}', 
            '{{DEPARTURE_DATE}}',
            '{{ORIGIN_AIRPORT_NAME}}',
            '{{DESTINATION_AIRPORT_NAME}}',
            '{{CREW_COUNT}}',
            '{{TOTAL_PAX_COUNT}}',
            '{{PRIMARY_PILOT}}',
            '{{PASSENGERS_AND_PATIENT_LIST}}'
        ]
        
        found_placeholders = []
        for placeholder in critical_placeholders:
            if placeholder in all_text:
                print(f" {placeholder}")
                found_placeholders.append(placeholder)
            else:
                print(f" {placeholder}")
        
        completion_rate = (len(found_placeholders) / len(critical_placeholders)) * 100
        print(f"\n Placeholder completion: {len(found_placeholders)}/{len(critical_placeholders)} ({completion_rate:.1f}%)")
        
        # Test 4: Document modification and placeholder replacement
        print(f"\n TEST 4: PLACEHOLDER REPLACEMENT")
        print("-" * 40)
        
        # Create test data
        test_data = {
            '{{AIRCRAFT_COMPANY}}': 'JET Charter Services LLC',
            '{{AIRCRAFT_TAIL}}': 'N123JT',
            '{{AIRCRAFT_MAKE}}': 'Cessna',
            '{{AIRCRAFT_MODEL}}': 'Citation X',
            '{{FLIGHT_NUMBER}}': 'JET001-TEST',
            '{{CARRIER_CODE}}': 'JCS',
            '{{DEPARTURE_DATE}}': '09/02/2024',
            '{{DEPARTURE_TIME_UTC}}': '14:30',
            '{{ARRIVAL_TIME_UTC}}': '19:45',
            '{{ORIGIN_AIRPORT_NAME}}': 'Los Angeles International',
            '{{ORIGIN_AIRPORT_ICAO}}': 'KLAX',
            '{{DESTINATION_AIRPORT_NAME}}': 'John F. Kennedy International',
            '{{DESTINATION_AIRPORT_ICAO}}': 'KJFK',
            '{{CREW_COUNT}}': '3',
            '{{TOTAL_PAX_COUNT}}': '2',
            '{{PRIMARY_PILOT}}': 'Captain John Smith',
            '{{SECONDARY_PILOT}}': 'First Officer Sarah Johnson',
            '{{PASSENGERS_AND_PATIENT_LIST}}': 'John Doe (Patient)\nJane Smith (Passenger)',
            '{{DOCUMENT_DATE}}': datetime.now().strftime('%m/%d/%Y')
        }
        
        # Apply replacements to document
        replacements_applied = 0
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            for placeholder, value in test_data.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
                    if paragraph.text != original_text:
                        replacements_applied += 1
                        original_text = paragraph.text
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        for placeholder, value in test_data.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
                                if paragraph.text != original_text:
                                    replacements_applied += 1
                                    original_text = paragraph.text
        
        print(f" Placeholder replacements applied: {replacements_applied}")
        
        # Test 5: Save filled document
        print(f"\n TEST 5: DOCUMENT GENERATION")
        print("-" * 40)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"GenDec_Standalone_Test_{timestamp}.docx"
        output_path = output_dir / output_filename
        
        # Save filled document
        doc.save(output_path)
        
        # Verify saved document
        saved_size = output_path.stat().st_size
        print(f" Document saved: {output_filename}")
        print(f" Output size: {saved_size:,} bytes")
        
        # Test 6: Verify filled document
        print(f"\n TEST 6: FILLED DOCUMENT VERIFICATION")
        print("-" * 40)
        
        # Load and verify the filled document
        filled_doc = Document(output_path)
        filled_text = ""
        
        for paragraph in filled_doc.paragraphs:
            filled_text += paragraph.text + "\n"
            
        for table in filled_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    filled_text += cell.text + " "
        
        remaining_placeholders = filled_text.count('{{')
        print(f" Remaining placeholders: {remaining_placeholders}")
        print(f" Contains test data:")
        
        test_verifications = [
            ('Flight Number', 'JET001-TEST' in filled_text),
            ('Aircraft Registration', 'N123JT' in filled_text),
            ('Captain Name', 'Captain John Smith' in filled_text),
            ('Passenger Info', 'John Doe (Patient)' in filled_text),
            ('Airport Codes', 'KLAX' in filled_text and 'KJFK' in filled_text),
            ('Company Name', 'JET Charter Services LLC' in filled_text)
        ]
        
        verified_count = 0
        for check_name, result in test_verifications:
            status = "" if result else ""
            print(f"   {status} {check_name}")
            if result:
                verified_count += 1
        
        verification_rate = (verified_count / len(test_verifications)) * 100
        
        # Test 7: CBP Form compliance
        print(f"\n TEST 7: CBP FORM 7507 COMPLIANCE")
        print("-" * 40)
        
        compliance_checks = [
            ('OMB Control Number', '1651-0002' in filled_text),
            ('CBP Form Number', '7507' in filled_text),
            ('DHS Header', 'DEPARTMENT OF HOMELAND SECURITY' in filled_text),
            ('General Declaration Title', 'GENERAL DECLARATION' in filled_text),
            ('CFR References', '19 CFR' in filled_text),
            ('Two Page Structure', 'Page 1 of 2' in filled_text and 'Page 2 of 2' in filled_text)
        ]
        
        compliant_count = 0
        for check_name, result in compliance_checks:
            status = "" if result else ""
            print(f"   {status} {check_name}")
            if result:
                compliant_count += 1
        
        compliance_rate = (compliant_count / len(compliance_checks)) * 100
        
        # Final assessment
        print(f"\n TEST RESULTS SUMMARY")
        print("=" * 40)
        print(f" Template Structure: {'PASS' if len(doc.paragraphs) > 20 else 'FAIL'}")
        print(f" Placeholder System: {'PASS' if completion_rate >= 80 else 'FAIL'}")  
        print(f" Data Replacement: {'PASS' if remaining_placeholders == 0 else 'FAIL'}")
        print(f" Content Verification: {'PASS' if verification_rate >= 80 else 'FAIL'}")
        print(f" CBP Compliance: {'PASS' if compliance_rate >= 80 else 'FAIL'}")
        print(f" Document Generation: {'PASS' if saved_size > 30000 else 'FAIL'}")
        
        overall_success = (
            len(doc.paragraphs) > 20 and
            completion_rate >= 80 and
            remaining_placeholders == 0 and
            verification_rate >= 80 and
            compliance_rate >= 80 and
            saved_size > 30000
        )
        
        print(f"\n{' OVERALL RESULT: SUCCESS!' if overall_success else '  OVERALL RESULT: NEEDS ATTENTION'}")
        
        if overall_success:
            print(f" The fixed GenDec.docx template is fully functional")
            print(f" Ready for production use with CBP Form 7507 compliance")
        
        return overall_success
        
    except Exception as e:
        print(f" Test failed with error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    import sys
    success = test_standalone_template()
    print(f"\n{' ALL TESTS PASSED' if success else ' TESTS FAILED'}")
    sys.exit(0 if success else 1)
