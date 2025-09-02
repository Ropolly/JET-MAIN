#!/usr/bin/env python

import os
import sys
from pathlib import Path
from datetime import datetime, timedelta

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent.parent))

def create_mock_data(is_pax_leg=True, trip_number="JT001", origin_code="LAX", dest_code="JFK"):

    from datetime import datetime, timedelta
    
    # Create passengers and patient list for PAX legs
    passengers_and_patient = []
    passengers_info = []
    patient_info = None
    total_pax_count = 0
    
    if is_pax_leg:
        # Patient data
        patient_info = {
            'name': 'Patient Mary Johnson',
            'type': 'Patient',
            'nationality': 'USA',
            'passport_number': 'P87654321',
            'date_of_birth': '1978-03-15'
        }
        passengers_and_patient.append(patient_info)
        total_pax_count += 1
        
        # Passenger data
        passenger_list = [
            {
                'name': 'Dr. Robert Wilson',
                'type': 'Passenger',
                'nationality': 'USA', 
                'passport_number': 'P11111111',
                'date_of_birth': '1985-07-22'
            },
            {
                'name': 'Nurse Sarah Davis',
                'type': 'Passenger',
                'nationality': 'CAN',
                'passport_number': 'C22222222',
                'date_of_birth': '1990-12-10'
            }
        ]
        passengers_and_patient.extend(passenger_list)
        passengers_info = passenger_list
        total_pax_count += len(passenger_list)
    
    base_data = {
        # Flight information  
        'trip_number': trip_number,
        'flight_type': "Passenger" if is_pax_leg else "Repositioning",
        'departure_date': '2025-09-03 08:07',
        'departure_time_utc': '2025-09-03 08:07 UTC',
        'arrival_date': '2025-09-03 13:37',
        'arrival_time_utc': '2025-09-03 13:37 UTC',
        'flight_time': '05:30',
        'distance': '2,475 NM',
        
        # Airport information
        'origin_airport_name': 'Los Angeles International Airport' if origin_code == 'LAX' else 'John F. Kennedy International Airport',
        'origin_airport_code': origin_code,
        'origin_airport_icao': 'KLAX' if origin_code == 'LAX' else 'KJFK',
        'origin_city': 'Los Angeles' if origin_code == 'LAX' else 'New York',
        'origin_country': 'USA',
        
        'destination_airport_name': 'John F. Kennedy International Airport' if dest_code == 'JFK' else 'Miami International Airport',
        'destination_airport_code': dest_code,
        'destination_airport_icao': 'KJFK' if dest_code == 'JFK' else 'KMIA',
        'destination_city': 'New York' if dest_code == 'JFK' else 'Miami',
        'destination_country': 'USA',
        
        # Aircraft information (nested dict as expected)
        'aircraft': {
            'tail_number': 'N123JT',
            'make': 'Gulfstream',
            'model': 'G650',
            'company': 'JET Aviation Services',
            'serial_number': 'SN12345',
            'mgtow': '45,000 lbs'
        },
        
        # Crew information (nested dict as expected)
        'crew': {
            'primary_pilot': 'Captain John Smith',
            'secondary_pilot': 'First Officer Jane Doe',
            'medics': ['Dr. Emergency Med']
        },
        
        # Passenger information
        'passengers_and_patient': passengers_and_patient,
        'passengers': passengers_info,
        'patient': patient_info,
        'passenger_count': len(passengers_info),
        'total_pax_count': total_pax_count,
        'is_pax_leg': is_pax_leg,
        
        # Document generation metadata
        'document_date': datetime.now().strftime('%Y-%m-%d'),
        'document_time': datetime.now().strftime('%H:%M:%S'),
        
        'alternate_airport': 'LGA',
        'flight_rules': 'IFR',
        'flight_level': 'FL410'
    }
    
    return base_data

def create_repositioning_data():

    # Create repositioning data directly with correct parameters
    return create_mock_data(
        is_pax_leg=False,
        trip_number='JT001',
        origin_code='JFK',
        dest_code='MIA'
    )

def generate_mock_documents():

    print("="*60)
    print("MOCK DOCUMENT GENERATION")
    print("="*60)
    
    # Create standalone document generator (no Django imports needed)
    try:
        class MockDocumentGenerator:

            def __init__(self):
                self.templates_dir = Path(__file__).parent / "documents"
                self.outputs_dir = Path(__file__).parent / "outputs"
                self.outputs_dir.mkdir(exist_ok=True)
            
            def _fill_docx_template(self, template_path, output_path, data):

                try:
                    # Try using python-docx
                    from docx import Document
                    
                    print(f" Using template: {template_path.name}")
                    print(f" Generating: {output_path.name}")
                    
                    # Load template
                    doc = Document(template_path)
                    
                    # Replace placeholders in paragraphs
                    for paragraph in doc.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, data)
                    
                    # Replace placeholders in tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self._replace_placeholders_in_paragraph(paragraph, data)
                    
                    # Handle special passenger list table
                    if data.get('passengers_and_patient'):
                        self._fill_passenger_table(doc, data)
                    
                    # Save document
                    doc.save(output_path)
                    print(f" Generated: {output_path}")
                    return str(output_path)
                    
                except ImportError:
                    print("  python-docx not available, creating simple text file")
                    return self._generate_text_fallback(template_path, output_path, data)
                except Exception as e:
                    print(f" Error with python-docx: {e}")
                    return self._generate_text_fallback(template_path, output_path, data)
            
            def _replace_placeholders_in_paragraph(self, paragraph, data):

                if not paragraph.text:
                    return
                
                text = paragraph.text
                
                # Create replacements dictionary matching real DocumentGenerator
                replacements = self._create_replacements_dict(data)
                
                # Replace all placeholders
                for placeholder, value in replacements.items():
                    if placeholder in text:
                        text = text.replace(placeholder, str(value))
                
                # Update paragraph if text changed
                if text != paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(text)
            
            def _create_replacements_dict(self, data):

                replacements = {
                    '{{TRIP_NUMBER}}': data['trip_number'],
                    '{{FLIGHT_TYPE}}': data['flight_type'],
                    '{{DEPARTURE_DATE}}': data['departure_date'],
                    '{{DEPARTURE_TIME_UTC}}': data['departure_time_utc'],
                    '{{ARRIVAL_DATE}}': data['arrival_date'],
                    '{{ARRIVAL_TIME_UTC}}': data['arrival_time_utc'],
                    '{{FLIGHT_TIME}}': data['flight_time'],
                    '{{DISTANCE}}': data['distance'],
                    
                    '{{ORIGIN_AIRPORT}}': data['origin_airport_name'],
                    '{{ORIGIN_CODE}}': data['origin_airport_code'],
                    '{{ORIGIN_ICAO}}': data['origin_airport_icao'],
                    '{{ORIGIN_CITY}}': data['origin_city'],
                    '{{ORIGIN_COUNTRY}}': data['origin_country'],
                    
                    '{{DESTINATION_AIRPORT}}': data['destination_airport_name'],
                    '{{DESTINATION_CODE}}': data['destination_airport_code'],
                    '{{DESTINATION_ICAO}}': data['destination_airport_icao'],
                    '{{DESTINATION_CITY}}': data['destination_city'],
                    '{{DESTINATION_COUNTRY}}': data['destination_country'],
                    
                    '{{AIRCRAFT_TAIL}}': data['aircraft']['tail_number'],
                    '{{AIRCRAFT_MAKE}}': data['aircraft']['make'],
                    '{{AIRCRAFT_MODEL}}': data['aircraft']['model'],
                    '{{AIRCRAFT_COMPANY}}': data['aircraft']['company'],
                    '{{AIRCRAFT_SERIAL}}': data['aircraft']['serial_number'],
                    '{{AIRCRAFT_MGTOW}}': data['aircraft']['mgtow'],
                    
                    '{{PRIMARY_PILOT}}': data['crew']['primary_pilot'],
                    '{{SECONDARY_PILOT}}': data['crew']['secondary_pilot'],
                    '{{MEDICS}}': ', '.join(data['crew']['medics']) if data['crew']['medics'] else 'N/A',
                    
                    '{{PASSENGER_COUNT}}': str(data['passenger_count']),
                    '{{TOTAL_PAX_COUNT}}': str(data['total_pax_count']),
                    '{{IS_PAX_LEG}}': 'Yes' if data['is_pax_leg'] else 'No',
                    '{{DOCUMENT_DATE}}': data['document_date'],
                    '{{DOCUMENT_TIME}}': data['document_time'],
                }
                
                # Add combined passenger and patient list (PAX leg logic)
                if data['passengers_and_patient']:
                    pax_list = '\n'.join([
                        f"{p['name']} ({p['type']}) - {p['nationality']} - Passport: {p['passport_number']}"
                        for p in data['passengers_and_patient']
                    ])
                    replacements['{{PASSENGERS_AND_PATIENT_LIST}}'] = pax_list
                else:
                    replacements['{{PASSENGERS_AND_PATIENT_LIST}}'] = 'No passengers or patient (Repositioning flight)'
                
                # Add legacy passenger list for backward compatibility
                if data['passengers']:
                    passenger_list = '\n'.join([
                        f"{p['name']} - {p['nationality']} - Passport: {p['passport_number']}"
                        for p in data['passengers']
                    ])
                    replacements['{{PASSENGER_LIST}}'] = passenger_list
                else:
                    replacements['{{PASSENGER_LIST}}'] = 'No passengers (Repositioning flight)'
                
                # Add patient info if available
                if data['patient']:
                    replacements['{{PATIENT_NAME}}'] = data['patient']['name']
                    replacements['{{PATIENT_DOB}}'] = data['patient']['date_of_birth']
                else:
                    replacements['{{PATIENT_NAME}}'] = 'N/A'
                    replacements['{{PATIENT_DOB}}'] = 'N/A'
                
                return replacements
            
            def _fill_passenger_table(self, doc, data):

                passengers = data.get('passengers_and_patient', [])
                if not passengers:
                    return
                
                # Look for tables that might be passenger tables
                for table in doc.tables:
                    if len(table.rows) > 1 and len(table.columns) >= 4:
                        # Check if this looks like a passenger table
                        header_text = ' '.join([cell.text.lower() for cell in table.rows[0].cells])
                        if any(word in header_text for word in ['name', 'passenger', 'type', 'nationality', 'passport']):
                            # Clear existing rows except header
                            for i in range(len(table.rows) - 1, 0, -1):
                                table._element.remove(table.rows[i]._element)
                            
                            # Add passenger rows
                            for passenger in passengers:
                                row = table.add_row()
                                cells = row.cells
                                if len(cells) >= 5:
                                    cells[0].text = passenger['name']
                                    cells[1].text = passenger['type']
                                    cells[2].text = passenger['nationality']
                                    cells[3].text = passenger['passport_number']
                                    cells[4].text = passenger['date_of_birth']
                            break
            
            def _format_passenger_list(self, passengers):

                if not passengers:
                    return "No passengers"
                
                lines = []
                for i, pax in enumerate(passengers, 1):
                    lines.append(f"{i}. {pax['name']} ({pax['type']}) - {pax['nationality']} - {pax['passport_number']}")
                return '\n'.join(lines)
            
            def _generate_text_fallback(self, template_path, output_path, data):

                output_path = output_path.with_suffix('.txt')
                
                content = f"""
MOCK DOCUMENT: {template_path.stem}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

FLIGHT INFORMATION:
Trip Number: {data.get('trip_number', 'N/A')}
Departure: {data.get('departure_time', 'N/A')}
Arrival: {data.get('arrival_time', 'N/A')}
Flight Time: {data.get('flight_time', 'N/A')}

ROUTE:
From: {data.get('origin_airport_name', 'N/A')} ({data.get('origin_airport_code', 'N/A')})
To: {data.get('destination_airport_name', 'N/A')} ({data.get('destination_airport_code', 'N/A')})

AIRCRAFT:
Tail Number: {data.get('aircraft_tail_number', 'N/A')}
Make/Model: {data.get('aircraft_make', 'N/A')} {data.get('aircraft_model', 'N/A')}
Company: {data.get('aircraft_company', 'N/A')}

CREW:
Captain: {data.get('primary_in_command', 'N/A')}
First Officer: {data.get('secondary_in_command', 'N/A')}

PASSENGERS AND PATIENT:
PAX Leg: {data.get('is_pax_leg', False)}
Total Count: {data.get('total_pax_count', 0)}
"""
                
                if data.get('passengers_and_patient'):
                    content += "\nPassenger List:\n"
                    for i, pax in enumerate(data['passengers_and_patient'], 1):
                        content += f"{i}. {pax['name']} ({pax['type']}) - {pax['nationality']} - {pax['passport_number']}\n"
                else:
                    content += "\nNo passengers/patient on this leg.\n"
                
                with open(output_path, 'w') as f:
                    f.write(content)
                
                print(f" Generated text fallback: {output_path}")
                return str(output_path)
        
        # Generate documents
        generator = MockDocumentGenerator()
        
        # Generate PAX leg documents
        print("\n--- Generating PAX Leg Documents ---")
        pax_data = create_mock_data()
        
        gendec_template = generator.templates_dir / "GenDec.docx"
        hr_template = generator.templates_dir / "HandlingRequest.docx"
        
        if gendec_template.exists():
            gendec_output = generator.outputs_dir / "Mock_GenDec_PAX_JT001_LAX_JFK.docx"
            generator._fill_docx_template(gendec_template, gendec_output, pax_data)
        else:
            print(f" GenDec template not found: {gendec_template}")
        
        if hr_template.exists():
            hr_output = generator.outputs_dir / "Mock_HandlingRequest_PAX_JT001_LAX_JFK.docx"
            generator._fill_docx_template(hr_template, hr_output, pax_data)
        else:
            print(f" HandlingRequest template not found: {hr_template}")
        
        # Generate repositioning leg documents
        print("\n--- Generating Repositioning Leg Documents ---")
        repo_data = create_repositioning_data()
        
        if gendec_template.exists():
            gendec_repo_output = generator.outputs_dir / "Mock_GenDec_REPO_JT001_JFK_MIA.docx"
            generator._fill_docx_template(gendec_template, gendec_repo_output, repo_data)
        
        if hr_template.exists():
            hr_repo_output = generator.outputs_dir / "Mock_HandlingRequest_REPO_JT001_JFK_MIA.docx"
            generator._fill_docx_template(hr_template, hr_repo_output, repo_data)
        
        print(f"\n Mock documents generated in: {generator.outputs_dir}")
        print(f" Check the outputs directory for generated files")
        
    except Exception as e:
        print(f" Error generating mock documents: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    generate_mock_documents()
