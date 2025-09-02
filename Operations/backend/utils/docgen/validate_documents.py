#!/usr/bin/env python

import os
from pathlib import Path
from docx import Document

def validate_docx_document(docx_path):

    print(f"\n Validating: {docx_path.name}")
    
    try:
        doc = Document(docx_path)
        
        # Extract all text from document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text.append(paragraph.text)
        
        document_text = '\n'.join(full_text)
        
        # Check for unreplaced placeholders
        unreplaced_placeholders = []
        for line in document_text.split('\n'):
            if '{{' in line and '}}' in line:
                # Extract placeholders
                import re
                placeholders = re.findall(r'\{\{[^}]+\}\}', line)
                unreplaced_placeholders.extend(placeholders)
        
        # Check for expected content
        expected_content = {
            'Trip Number': 'JT001' in document_text,
            'Aircraft Tail': 'N123JT' in document_text,
            'Origin Airport': 'LAX' in document_text or 'Los Angeles' in document_text,
            'Destination Airport': 'JFK' in document_text or 'John F. Kennedy' in document_text,
            'Aircraft Make': 'Gulfstream' in document_text,
            'Captain Name': 'Captain John Smith' in document_text,
        }
        
        # Check PAX specific content for PAX documents
        if 'PAX' in docx_path.name:
            expected_content.update({
                'Patient Name': 'Patient Mary Johnson' in document_text or 'Mary Johnson' in document_text,
                'Passenger Name': 'Dr. Robert Wilson' in document_text or 'Robert Wilson' in document_text,
                'Nationality': 'USA' in document_text,
                'Passport': 'P87654321' in document_text or 'P11111111' in document_text,
            })
        
        # Report results
        print(f"  Unreplaced placeholders: {len(unreplaced_placeholders)}")
        if unreplaced_placeholders:
            for placeholder in set(unreplaced_placeholders):
                print(f"    - {placeholder}")
        
        print(f"  Expected content check:")
        all_good = True
        for content_type, found in expected_content.items():
            status = "" if found else ""
            print(f"    {status} {content_type}")
            if not found:
                all_good = False
        
        if all_good and not unreplaced_placeholders:
            print(f"   Document validation: PASSED")
        else:
            print(f"    Document validation: NEEDS ATTENTION")
            
        return len(unreplaced_placeholders), all_good
        
    except Exception as e:
        print(f"   Error validating document: {e}")
        return -1, False

def main():
    print("="*60)
    print("DOCUMENT VALIDATION")
    print("="*60)
    
    outputs_dir = Path(__file__).parent / "outputs"
    
    if not outputs_dir.exists():
        print(" Outputs directory not found!")
        return
    
    # Find all .docx files
    docx_files = list(outputs_dir.glob("*.docx"))
    
    if not docx_files:
        print(" No .docx files found in outputs directory!")
        return
    
    print(f"Found {len(docx_files)} Word documents to validate")
    
    total_placeholders = 0
    total_passed = 0
    
    for docx_file in docx_files:
        placeholder_count, passed = validate_docx_document(docx_file)
        if placeholder_count >= 0:
            total_placeholders += placeholder_count
            if passed:
                total_passed += 1
    
    print("\n" + "="*60)
    print("VALIDATION SUMMARY")
    print("="*60)
    print(f"Documents validated: {len(docx_files)}")
    print(f"Documents passed: {total_passed}")
    print(f"Total unreplaced placeholders: {total_placeholders}")
    
    if total_passed == len(docx_files) and total_placeholders == 0:
        print(" All documents validated successfully!")
    else:
        print("  Some documents need attention")

if __name__ == "__main__":
    main()
