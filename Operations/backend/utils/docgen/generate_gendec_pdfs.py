#!/usr/bin/env python3

from pathlib import Path
import shutil
import sys
from datetime import datetime
from docx import Document

def create_standalone_generator():

    class StandaloneDocumentGenerator:
        def __init__(self):
            self.base_dir = Path(__file__).parent
            self.templates_dir = self.base_dir / "documents"
            self.outputs_dir = self.base_dir / "pdf_outputs"
            self.outputs_dir.mkdir(exist_ok=True)
        
        def generate_gendec_pdf(self, data):

            # Use existing GenDec.docx template
            template_path = self.templates_dir / "GenDec.docx"
            if not template_path.exists():
                raise FileNotFoundError(f"GenDec template not found: {template_path}")
            
            # Create output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            flight_type = "PAX" if data.get('is_pax_leg', False) else "REPO"
            flight_num = data.get('flight_number', 'UNKNOWN')
            output_filename = f"GenDec_{flight_type}_{flight_num}_{timestamp}.docx"
            output_path = self.outputs_dir / output_filename
            
            # Fill template with data
            self._fill_docx_template(template_path, output_path, data)
            
            # Convert to PDF (simulate)
            pdf_filename = output_filename.replace('.docx', '.pdf')
            pdf_path = self.outputs_dir / pdf_filename
            
            # Create a text file representing the PDF content
            self._create_pdf_representation(pdf_path, data)
            
            return str(pdf_path)
        
        def _fill_docx_template(self, template_path, output_path, data):

            # Load template
            doc = Document(template_path)
            
            # Prepare replacements
            replacements = self._create_replacements(data)
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Save filled document
            doc.save(output_path)
            print(f" GenDec Word document generated: {output_path.name}")
        
        def _create_replacements(self, data):

            # Format passengers and patient list
            passenger_list = ""
            if data.get('passengers_and_patient'):
                lines = []
                for i, person in enumerate(data['passengers_and_patient'], 1):
                    line = f"{i}. {person['name']} ({person['type']}) - {person['nationality']} - Passport: {person['passport_number']}"
                    lines.append(line)
                passenger_list = '\n'.join(lines)
            else:
                passenger_list = "No passengers or patient (Repositioning flight)"
            
            # Create replacements dictionary
            replacements = {
                '{{TRIP_NUMBER}}': data.get('trip_number', 'N/A'),
                '{{DEPARTURE_DATE}}': data.get('departure_date', 'N/A'),
                '{{DEPARTURE_TIME_UTC}}': data.get('departure_time_utc', 'N/A'),
                '{{ARRIVAL_DATE}}': data.get('arrival_date', 'N/A'),
                '{{ARRIVAL_TIME_UTC}}': data.get('arrival_time_utc', 'N/A'),
                '{{ORIGIN_AIRPORT_NAME}}': data.get('origin_airport_name', 'N/A'),
                '{{ORIGIN_AIRPORT_CODE}}': data.get('origin_airport_code', 'N/A'),
                '{{ORIGIN_AIRPORT_ICAO}}': data.get('origin_airport_icao', 'N/A'),
                '{{ORIGIN_CITY}}': data.get('origin_city', 'N/A'),
                '{{DESTINATION_AIRPORT_NAME}}': data.get('destination_airport_name', 'N/A'),
                '{{DESTINATION_AIRPORT_CODE}}': data.get('destination_airport_code', 'N/A'),
                '{{DESTINATION_AIRPORT_ICAO}}': data.get('destination_airport_icao', 'N/A'),
                '{{DESTINATION_CITY}}': data.get('destination_city', 'N/A'),
                '{{AIRCRAFT_TAIL}}': data.get('aircraft', {}).get('tail_number', 'N/A'),
                '{{AIRCRAFT_MAKE}}': data.get('aircraft', {}).get('make', 'N/A'),
                '{{AIRCRAFT_MODEL}}': data.get('aircraft', {}).get('model', 'N/A'),
                '{{AIRCRAFT_COMPANY}}': data.get('aircraft', {}).get('company', 'N/A'),
                '{{AIRCRAFT_SERIAL}}': data.get('aircraft', {}).get('serial_number', 'N/A'),
                '{{AIRCRAFT_MGTOW}}': data.get('aircraft', {}).get('mgtow', 'N/A'),
                '{{PRIMARY_PILOT}}': data.get('crew', {}).get('primary_pilot', 'N/A'),
                '{{SECONDARY_PILOT}}': data.get('crew', {}).get('secondary_pilot', 'N/A'),
                '{{CREW_COUNT}}': str(data.get('crew_count', 0)),
                '{{PASSENGER_COUNT}}': str(data.get('passenger_count', 0)),
                '{{TOTAL_PAX_COUNT}}': str(data.get('total_pax_count', 0)),
                '{{CARRIER_CODE}}': data.get('carrier_code', 'N/A'),
                '{{FLIGHT_NUMBER}}': data.get('flight_number', data.get('trip_number', 'N/A')),
                '{{PILOT_TITLE}}': data.get('pilot_title', 'Aircraft Commander'),
                '{{PASSENGERS_AND_PATIENT_LIST}}': passenger_list,
                '{{DOCUMENT_DATE}}': data.get('document_date', datetime.now().strftime('%Y-%m-%d')),
                '{{DOCUMENT_TIME}}': data.get('document_time', datetime.now().strftime('%H:%M:%S UTC')),
            }
            
            return replacements
        
        def _replace_placeholders_in_paragraph(self, paragraph, replacements):

            full_text = paragraph.text
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
            
            if full_text != paragraph.text:
                paragraph.text = full_text
        
        def _create_pdf_representation(self, pdf_path, data):

            with open(pdf_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("GENERATED GENDEC PDF CONTENT PREVIEW\n")
                f.write(f"Original Template: GenDec.docx\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("=" * 80 + "\n\n")
                
                # Flight information
                f.write("FLIGHT INFORMATION:\n")
                f.write("-" * 20 + "\n")
                f.write(f"Trip Number: {data.get('trip_number', 'N/A')}\n")
                f.write(f"Flight Number: {data.get('flight_number', 'N/A')}\n")
                f.write(f"Flight Type: {'PAX (Medical Charter)' if data.get('is_pax_leg') else 'Repositioning'}\n")
                f.write(f"Date: {data.get('departure_date', 'N/A')}\n")
                f.write(f"Departure: {data.get('departure_time_utc', 'N/A')} UTC\n")
                f.write(f"Arrival: {data.get('arrival_time_utc', 'N/A')} UTC\n\n")
                
                # Route information
                f.write("ROUTE:\n")
                f.write("-" * 10 + "\n")
                f.write(f"From: {data.get('origin_airport_name', 'N/A')} ({data.get('origin_airport_icao', 'N/A')})\n")
                f.write(f"      {data.get('origin_city', 'N/A')}\n")
                f.write(f"To:   {data.get('destination_airport_name', 'N/A')} ({data.get('destination_airport_icao', 'N/A')})\n")
                f.write(f"      {data.get('destination_city', 'N/A')}\n\n")
                
                # Aircraft information
                aircraft = data.get('aircraft', {})
                f.write("AIRCRAFT:\n")
                f.write("-" * 10 + "\n")
                f.write(f"Registration: {aircraft.get('tail_number', 'N/A')}\n")
                f.write(f"Make/Model: {aircraft.get('make', 'N/A')} {aircraft.get('model', 'N/A')}\n")
                f.write(f"Operator: {aircraft.get('company', 'N/A')}\n")
                f.write(f"Serial Number: {aircraft.get('serial_number', 'N/A')}\n")
                f.write(f"MGTOW: {aircraft.get('mgtow', 'N/A')} lbs\n\n")
                
                # Crew information
                crew = data.get('crew', {})
                f.write("CREW:\n")
                f.write("-" * 10 + "\n")
                f.write(f"Primary Pilot: {crew.get('primary_pilot', 'N/A')}\n")
                f.write(f"Secondary Pilot: {crew.get('secondary_pilot', 'N/A')}\n")
                f.write(f"Total Crew: {data.get('crew_count', 0)}\n")
                
                medics = crew.get('medics', [])
                if medics:
                    f.write("Medical Crew:\n")
                    for medic in medics:
                        f.write(f"  • {medic}\n")
                f.write("\n")
                
                # Passenger information
                f.write("PASSENGERS:\n")
                f.write("-" * 15 + "\n")
                f.write(f"Total Passengers: {data.get('total_pax_count', 0)}\n")
                
                if data.get('passengers_and_patient'):
                    f.write("\nManifest:\n")
                    for i, person in enumerate(data['passengers_and_patient'], 1):
                        f.write(f"{i}. {person['name']} ({person['type']})\n")
                        f.write(f"   Nationality: {person['nationality']}\n")
                        f.write(f"   Passport: {person['passport_number']}\n")
                        f.write(f"   DOB: {person['date_of_birth']}\n\n")
                else:
                    f.write("No passengers (Repositioning flight)\n\n")
                
                f.write("=" * 80 + "\n")
                f.write("This represents the content that would be in the actual PDF file\n")
                f.write("generated from the GenDec.docx template with filled data.\n")
                f.write("=" * 80 + "\n")
    
    return StandaloneDocumentGenerator()

def create_mock_pax_data():

    return {
        'trip_number': 'JET-2024-0902',
        'departure_date': '09/02/2024',
        'departure_time_utc': '14:30:00',
        'arrival_date': '09/02/2024',
        'arrival_time_utc': '18:45:00',
        
        'origin_airport_name': 'Teterboro Airport',
        'origin_airport_code': 'TEB',
        'origin_airport_icao': 'KTEB',
        'origin_city': 'Teterboro, NJ',
        
        'destination_airport_name': 'Miami International Airport',
        'destination_airport_code': 'MIA',
        'destination_airport_icao': 'KMIA',
        'destination_city': 'Miami, FL',
        
        'aircraft': {
            'tail_number': 'N525JT',
            'make': 'Cessna',
            'model': 'Citation CJ3+',
            'company': 'JET Charter Services LLC',
            'serial_number': '525C-0924',
            'mgtow': '13,870'
        },
        
        'crew': {
            'primary_pilot': 'Captain John Smith',
            'secondary_pilot': 'First Officer Sarah Johnson',
            'medics': ['Dr. Michael Brown', 'Paramedic Lisa Davis']
        },
        
        'passengers_and_patient': [
            {
                'name': 'Robert Wilson',
                'type': 'Patient',
                'nationality': 'United States',
                'passport_number': 'P123456789',
                'date_of_birth': '1965-03-15'
            },
            {
                'name': 'Emily Wilson',
                'type': 'Passenger',
                'nationality': 'United States',
                'passport_number': 'P987654321',
                'date_of_birth': '1968-07-22'
            },
            {
                'name': 'Dr. Jennifer Thompson',
                'type': 'Passenger',
                'nationality': 'United States',
                'passport_number': 'P456789123',
                'date_of_birth': '1972-11-08'
            }
        ],
        
        'passenger_count': 2,
        'total_pax_count': 3,
        'is_pax_leg': True,
        'crew_count': 4,
        'carrier_code': 'JCS',
        'flight_number': 'JET001',
        'pilot_title': 'Aircraft Commander'
    }

def create_mock_repo_data():

    return {
        'trip_number': 'JET-2024-0902R',
        'departure_date': '09/02/2024',
        'departure_time_utc': '20:15:00',
        'arrival_date': '09/02/2024',
        'arrival_time_utc': '20:45:00',
        
        'origin_airport_name': 'Miami International Airport',
        'origin_airport_code': 'MIA',
        'origin_airport_icao': 'KMIA',
        'origin_city': 'Miami, FL',
        
        'destination_airport_name': 'Fort Lauderdale Executive Airport',
        'destination_airport_code': 'FXE',
        'destination_airport_icao': 'KFXE',
        'destination_city': 'Fort Lauderdale, FL',
        
        'aircraft': {
            'tail_number': 'N525JT',
            'make': 'Cessna',
            'model': 'Citation CJ3+',
            'company': 'JET Charter Services LLC',
            'serial_number': '525C-0924',
            'mgtow': '13,870'
        },
        
        'crew': {
            'primary_pilot': 'Captain John Smith',
            'secondary_pilot': 'First Officer Sarah Johnson',
            'medics': []
        },
        
        'passengers_and_patient': [],
        
        'passenger_count': 0,
        'total_pax_count': 0,
        'is_pax_leg': False,
        'crew_count': 2,
        'carrier_code': 'JCS',
        'flight_number': 'JET001R',
        'pilot_title': 'Aircraft Commander'
    }

def main():

    print(" GENERATING GENDEC PDF OUTPUTS")
    print("Using existing GenDec.docx template")
    print("=" * 60)
    
    try:
        # Create generator
        generator = create_standalone_generator()
        
        # Test PAX leg
        print(" Generating PAX Leg GenDec...")
        pax_data = create_mock_pax_data()
        pax_pdf = generator.generate_gendec_pdf(pax_data)
        print(f" PAX leg generated: {Path(pax_pdf).name}")
        
        # Test repositioning leg
        print("\n Generating Repositioning Leg GenDec...")
        repo_data = create_mock_repo_data()
        repo_pdf = generator.generate_gendec_pdf(repo_data)
        print(f" Repositioning leg generated: {Path(repo_pdf).name}")
        
        # Summary
        output_dir = generator.outputs_dir
        print(f"\n Generated files in: {output_dir}")
        
        # List generated files
        docx_files = list(output_dir.glob("*.docx"))
        txt_files = list(output_dir.glob("*.txt"))
        
        print(f"\n DOCX Files (filled templates):")
        for f in docx_files:
            size_kb = f.stat().st_size / 1024
            print(f"   • {f.name} ({size_kb:.1f} KB)")
        
        print(f"\n TXT Files (PDF content previews):")
        for f in txt_files:
            size_kb = f.stat().st_size / 1024
            print(f"   • {f.name} ({size_kb:.1f} KB)")
        
        print(f"\n GenDec PDF generation complete!")
        print(f" Review the generated DOCX files to verify template filling")
        print(f" Review the TXT files to preview PDF content")
        
    except Exception as e:
        print(f" Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
