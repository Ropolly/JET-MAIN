#!/usr/bin/env python
"""
Script to check template content for placeholders
"""

import os
import sys
import django
from pathlib import Path

# Setup Django environment
sys.path.append(str(Path(__file__).parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'backend.settings')
django.setup()

try:
    from docx import Document as DocxDocument
    print("python-docx is available")
except ImportError:
    print("python-docx not available")
    exit(1)

def check_template_content():
    """Check what placeholders exist in template files"""
    templates_dir = Path("utils/docgen/documents")
    
    if not templates_dir.exists():
        print("Templates directory not found")
        return
    
    template_files = ["CustomerItin.docx", "HandlingRequest.docx", "GenDec.docx"]
    
    for template_name in template_files:
        template_path = templates_dir / template_name
        if not template_path.exists():
            print(f"Template not found: {template_name}")
            continue
            
        print(f"\n=== Checking {template_name} ===")
        try:
            doc = DocxDocument(template_path)
            all_text = ""
            
            # Extract all text from document
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + " "
            
            # Check for tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
            
            print(f"Total characters in template: {len(all_text)}")
            
            # Look for common placeholders
            common_placeholders = [
                '{{TRIP_NUMBER}}', '{{AIRCRAFT_TAIL}}', '{{PRIMARY_PILOT}}', 
                '{{DEPARTURE_DATE}}', '{{ORIGIN_AIRPORT}}', '{{DESTINATION_AIRPORT}}',
                '{{COMPANY_NAME}}', '{{FLIGHT_TYPE}}', '{{PASSENGER_COUNT}}'
            ]
            
            found_placeholders = []
            for placeholder in common_placeholders:
                if placeholder in all_text:
                    found_placeholders.append(placeholder)
            
            if found_placeholders:
                print(f"Found placeholders: {found_placeholders}")
            else:
                print("No standard placeholders found")
                # Show a sample of the content
                sample = all_text[:200].strip()
                print(f"Sample content: {sample}...")
                
                # Look for any {{ }} patterns
                import re
                placeholder_patterns = re.findall(r'\{\{[^}]+\}\}', all_text)
                if placeholder_patterns:
                    print(f"Found other placeholder patterns: {placeholder_patterns[:10]}")  # Show first 10
                else:
                    print("No placeholder patterns ({{ }}) found at all")
                    
        except Exception as e:
            print(f"Error reading {template_name}: {e}")

if __name__ == "__main__":
    check_template_content()